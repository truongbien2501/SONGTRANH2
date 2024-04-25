from pydoc import doc
from docx import Document
from docx.shared import Pt,Inches
import os
from win32com import client
# # WORD 
# word = client.Dispatch("Word.Application")
# word.Visible = True

# pth = r'\\admin-pc\DATA tin\TRUYENHINH\GMDB-KKL\2022'
# lisd = os.listdir(pth)

# for a in lisd:
#     if 'docx' not in a  and 'doc' in a and '$' not in a:
#         print(a)
#         odoc = word.Documents.Open(pth + '/' + a)
#         odoc.SaveAs(pth + '/' + a + 'x', FileFormat=16)
#         odoc.Close()

# word.Quit()



# def ngaythangphathanh(tentin):
#     a = tentin.spilit('_')

# pth = 'TINMAU/DRHV05.docx'
# pth = r'C:\Users\Admin\Desktop\Chuyengiao\testtin\QBIN_TVHN_20230814_1130.docx'
pth = 'TINMAU/ST_LULU.docx'
hsm = Document(pth)
i=-1
j =-1
for a in hsm.tables[4].rows:
    i=i+1
    j=-1
    for b in a.cells:
        j+=1
        print(b.text,i,j)
# print(hsm.tables[0].rows[1].cells[0].text)

# print(hsm.tables[0].cell(0, 14).text)
# if hsm.tables[3].cell(0,0).e:
#     print(hsm.tables[0].cell(0,0).text)

# 'Thời điểm dự báo'



# pth = r'\\admin-pc\DATA tin\TRUYENHINH\CANH BAO MUA DONG\QIII'
# list = os.listdir(pth)
# # print(list)
# for a in list[:1]:
#     if 'docx' in a or 'doc' in a:
#         tt = a.split("_")
#         nam = tt[2][:4]
#         thang = tt[2][4:6]
#         ngay = tt[2][6:8]
#         gio = tt[3][:2]
#         phut = tt[3][2:4]
#         hsm.paragraphs[1].text = 'Thời gian phát tin:' + gio + 'h' + phut + ' Ngày ' + ngay + ' tháng ' + thang + ' năm ' + nam

#         odoc = Document(pth + '/' + a)
#         for i in odoc.paragraphs:
#             if 'Bình Sơn' in i.text:
#                 hsm.tables[0].cell(3,10).paragraphs[1].text = 'Bình Sơn'
#         hsm.save('bantest.docx')
#         os.system('start bantest.docx')

