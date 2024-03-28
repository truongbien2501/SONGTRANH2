from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
from func.load_data import mo_word
selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value


def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def sobt():
    pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    # print(pth)
    now = datetime.now()
    if now.strftime('%Y%m%d') in pth:
        os.remove(pth)
        messagebox.showinfo('Thong bao','Đã xóa file tồn tại' + pth.split('\\')[-1])
        pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    
    odoc = Document(pth)
    for a in odoc.tables[0].cell(0,0).paragraphs:
        if 'Số' in a.text:
            dl = str(a.text)
            sbt = dl[dl.index('-')+1:dl.index('/')]
    return int(sbt) + 1

def xacdinhngaydb():
    now = datetime.now()
    for a in range(3,10):
        tttt = now + timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def xacdinhngaydaqua():
    now = datetime.now()
    for a in range(3,10):
        tttt = now - timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def tin_nenKT_05day():
    now = datetime.now()
    odoc = Document('TINMAU/DRHV05.docx')
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    # so ban tin
    sbt = sobt()
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số:ĐRHV-'+ str(sbt) + '/QNGA'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Ngãi' in dl:
                pr.text=''
                ntn = 'Quảng Ngãi, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
            
    
    ngaydb = xacdinhngaydb()
    ngaydb = ngaydb - timedelta(days=1)
    ngaydb = ngaydb.strftime('%d/%m/%Y')
    ngaytd = xacdinhngaydaqua()
    ngaytd = ngaytd.strftime('%d/%m/%Y')
    
    
    for pr in odoc.paragraphs:
        dl = pr.text
        if '(Từ ngày 11/9/2023 đến ngày 15/9/2023)' in dl:
            # ban tin tiep theo
            ntn = '(Từ ngày {} đến ngày {})'.format(now.strftime('%d/%m/%Y'),ngaydb)
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '* Đặc trưng thời tiết từ ngày 06/09 - 10/09/2023' in dl:
            # ban tin tiep theo
            ntn = '* Đặc trưng thời tiết từ ngày {} - {}'.format(ngaytd,(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '* Đặc trưng thủy văn từ ngày 06/9 - 10/9/2023' in dl:
            # ban tin tiep theo
            ntn = '* Đặc trưng thủy văn từ ngày {} - {}'.format(ngaytd,(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        elif '* Dự báo đặc trưng thời tiết 5 ngày tới' in dl:
            # ban tin tiep theo
            ntn = '* Dự báo đặc trưng thời tiết từ ngày {} đến ngày {}'.format(now.strftime('%d/%m/%Y'),ngaydb)
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13) 
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER      
        elif '* Dự báo đặc trưng thủy văn 5 ngày tới' in dl:
            # ban tin tiep theo
            ntn = '* Dự báo thuỷ văn từ ngày {} đến ngày {}'.format(now.strftime('%d/%m/%Y'),ngaydb)
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Trong 5 ngày tới (11/09 - 15/09/2023), lưu lượng nước về hồ' in dl:
            # ban tin tiep theo
            ntn = 'Trong 5 ngày tới ({} - {}), lưu lượng nước về hồ khả năng có biến động nhỏ vào chiều và tối. Đặc trưng lưu lượng và tổng lượng nước về hồ như sau:'.format(now.strftime('%d/%m/%Y'),ngaydb)
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'II/ DỰ BÁO 5 NGÀY TỚI (TỪ NGÀY 11/09/2023 ĐẾN NGÀY 15/09/2023)' in dl:
            # ban tin tiep theo
            ntn = 'II/ DỰ BÁO 5 NGÀY TỚI (TỪ NGÀY {} ĐẾN NGÀY {})'.format(now.strftime('%d/%m/%Y'),ngaydb)
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
                        
    # lay so lieu mua
    df= pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx',sheet_name='Muangay_theotin')
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    now = datetime(now.year,now.month,now.day,7)
    ngaylaymua = datetime.strptime(ngaytd, '%d/%m/%Y')
    # print(ngaytd)
    # print(ngaylaymua)
    df = df[(df['time'] <= now) & (df['time'] >= datetime(ngaylaymua.year,ngaylaymua.month,ngaylaymua.day,13))]
    df.set_index('time',inplace=True)
    mua6h = df.rolling(4,min_periods=1).sum()
    mua6h = mua6h.loc[mua6h.index.hour==7]
    mua6h = mua6h.applymap("{0:.1f}".format)
    # max_rain_dates = df.idxmax()
    max_values = df.max()
    tong_values = df.sum()
    for r in range(0,11):
        odoc.tables[1].cell(3,r).text = ''
        
    odoc.tables[1].cell(3,0).paragraphs[0].add_run("{0:.1f}".format(tong_values['Đầu mối'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,1).paragraphs[0].add_run(str(max_values['Đầu mối'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,2).paragraphs[0].add_run("{0:.1f}".format(tong_values['Đăk Nên'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,3).paragraphs[0].add_run(str(max_values['Đăk Nên'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,4).paragraphs[0].add_run("{0:.1f}".format(tong_values['Đăk tăng'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,5).paragraphs[0].add_run(str(max_values['Đăk tăng'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,6).paragraphs[0].add_run("{0:.1f}".format(tong_values['Sơn Tây'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,7).paragraphs[0].add_run(str(max_values['Sơn Tây'])).font.size = Pt(13) 
    
    # nhiet
    df= pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx',sheet_name='nhiet_am')
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="D")
    df['time'] = dt_rang
    now = datetime(now.year,now.month,now.day)
    ngaylaynhiet = datetime.strptime(ngaytd, '%d/%m/%Y')
    df = df[(df['time'] <= now) & (df['time'] > ngaylaynhiet)]
    df.set_index('time',inplace=True)
    
    max = df.agg(['mean','max','min'])

    odoc.tables[1].cell(3,8).paragraphs[0].add_run(str(max['nhiet_tb']['mean'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,9).paragraphs[0].add_run(str(max['nhiet_max']['max'])).font.size = Pt(13) 
    odoc.tables[1].cell(3,10).paragraphs[0].add_run(str(max['nhiet_min']['min'])).font.size = Pt(13) 
    
    
    
    # dem so ngay de them rows
    ngaydb = xacdinhngaydb()
    num_days = abs(ngaydb - now).days
    if num_days ==6:
        for a in range(3,5):
            odoc.tables[a].add_row()
    
    
    odoc.tables[3].cell(2,0).text = now.strftime('%d/%m')
    odoc.tables[3].cell(3,0).text = (now + timedelta(days=1)).strftime('%d/%m')
    odoc.tables[3].cell(4,0).text = (now + timedelta(days=2)).strftime('%d/%m')
    odoc.tables[3].cell(5,0).text = (now + timedelta(days=3)).strftime('%d/%m')
    odoc.tables[3].cell(6,0).text = (now + timedelta(days=4)).strftime('%d/%m')
    

    odoc.tables[4].cell(1,0).text = now.strftime('%d/%m')
    odoc.tables[4].cell(2,0).text = (now + timedelta(days=1)).strftime('%d/%m')
    odoc.tables[4].cell(3,0).text = (now + timedelta(days=2)).strftime('%d/%m')
    odoc.tables[4].cell(4,0).text = (now + timedelta(days=3)).strftime('%d/%m')
    odoc.tables[4].cell(5,0).text = (now + timedelta(days=4)).strftime('%d/%m')
            
    if num_days == 6:
        odoc.tables[3].cell(7,0).text = (now + timedelta(days=5)).strftime('%d/%m')
        odoc.tables[4].cell(6,0).text = (now + timedelta(days=5)).strftime('%d/%m')
    
 
    # q thuc do
    # odoc.tables[2].cell(3,10).paragraphs[0].add_run(str(max['nhiet_min']['min'])).font.size = Pt(13) 

    # print(mua6h)
    pth = read_txt('path_tin/DRHV.txt') + '/DHC_TV05_' + now.strftime('%Y%m%d_1600') + '.docx'
    odoc.save(pth)
    # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
def tin_tv05_load():
    now = datetime.now()
    pth = read_txt('path_tin/DRHV.txt') + '/DHC_TV05_' + now.strftime('%Y%m%d_1600') + '.docx'
    odoc = Document(pth)
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
   
    df = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',engine="openpyxl",sheet_name='DRHV05')
    
    df["Qtb_td"] = df["Qtb_td"].apply(lambda x: f"{x}")
    df["Qtb_td"] = df["Qtb_td"].astype(float)
    df["Qtb_td"] = df["Qtb_td"].map('{0:.1f}'.format)
    df["W"] = df["W"].apply(lambda x: f"{x}")
    df["W"] = df["W"].astype(float)
    df["W"] = df["W"].map('{0:.1f}'.format)
    
    df["Qtháng"] = df["Qtháng"].apply(lambda x: f"{x}")
    df["Qtháng"] = df["Qtháng"].astype(str)
    # df["Qtháng"] = df["Qtháng"].map('{0:.1f}'.format)
    df["W tháng"] = df["W tháng"].apply(lambda x: f"{x}")
    # df["Ngày"] = df["Ngày"].apply(lambda x: f"{x}")
    df["W tháng"] = df["W tháng"].astype(str)
    # df["W tháng"] = df["W tháng"].map('{0:.1f}'.format)
    # print(df.columns)
    
    
    now = datetime(now.year,now.month,now.day)
    kt = xacdinhngaydb()
    bd = xacdinhngaydaqua()

    df['time'] = pd.date_range(start=datetime(now.year,9,1), periods=len(df['Ngày']), freq="D")
    qtb = df[df['time']==now]['Qtb_td']
    tongluong = df[df['time']==now]['W']
    odoc.tables[2].cell(0,1).text = qtb
    odoc.tables[2].cell(1,1).text = tongluong
    
    
    df_db = df.loc[(df['time'] >= bd) & (df['time'] < now) ] # so lieu thuc 

    
    for i in range(1,6):  
        for j in range(1,3):         
            odoc.tables[4].cell(i,j).text = ''
    
    
    
    odoc.tables[4].cell(1,1).text = df[df['time']==now]['Qtháng']
    odoc.tables[4].cell(2,1).text = df[df['time']==(now + timedelta(days=1))]['Qtháng']
    odoc.tables[4].cell(3,1).text = df[df['time']==(now + timedelta(days=2))]['Qtháng']
    odoc.tables[4].cell(4,1).text = df[df['time']==(now + timedelta(days=3))]['Qtháng']
    odoc.tables[4].cell(5,1).text = df[df['time']==(now + timedelta(days=4))]['Qtháng']
    
    odoc.tables[4].cell(1,2).text = df[df['time']==now]['W tháng']
    odoc.tables[4].cell(2,2).text = df[df['time']==(now + timedelta(days=1))]['W tháng']
    odoc.tables[4].cell(3,2).text = df[df['time']==(now + timedelta(days=2))]['W tháng']
    odoc.tables[4].cell(4,2).text = df[df['time']==(now + timedelta(days=3))]['W tháng']
    odoc.tables[4].cell(5,2).text = df[df['time']==(now + timedelta(days=4))]['W tháng']
    
    ngaydb = xacdinhngaydb()
    num_days = abs(ngaydb - now).days
    if num_days ==6:
        odoc.tables[4].cell(6,1).text = df[df['time']==(now + timedelta(days=5))]['Qtháng']
        odoc.tables[4].cell(6,2).text = df[df['time']==(now + timedelta(days=5))]['W tháng']
    
    for i in range(1,6):  
        for j in range(0,3):         
            pr = odoc.tables[4].cell(i,j).paragraphs[0]
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    odoc.save(pth)
    mo_word(pth)
    # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
    