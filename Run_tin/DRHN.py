from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
import pyodbc
from func.CDH_TTB_API import TTB_API_mucnuoc
import numpy as np
import  mysql.connector
def creat_cxn():
    # Kết nối đến MySQL
    host = '113.160.225.84'
    user = 'qltram'
    password = 'mhq@123456'
    port = 3306
    database = 'datasolieu'
    cnx = mysql.connector.connect(host=host, user=user, password=password, port=port, database=database)
    return cnx

def query_sql(list_import,table_clounms,table_name):#TidVerticalIDVelocityForDetailMeasurement 
    if  str(list_import[2])=='nan':
        return None
    else:
        sql = 'INSERT INTO ' + table_name + '('
        gt =  " VALUES ("
        for a in range(len(table_clounms)):
            # print(table_clounms[a])
            # print(list_import[a])
            if str(list_import[a]) != 'nan':
                sql = sql + table_clounms[a]+ ','
                gt = gt + ',\'{}\''.format(list_import[a])
        sql = sql  + ')' + gt + ')'
        sql = sql.replace(',)',')')
        sql = sql.replace('(,','(')
        return sql


def insert_data(df,table_name):
    df.insert(0,'Matram','5ST')
    df['sldungduoc'] = df[df.columns[2]]
    df['maloi'] = 0
    df['chinhly'] = 0
    df = df.sort_values(by='time')
    # df =df.replace(np.nan,None)
    
    # Tạo kết nối
    cnx = creat_cxn()
    # Tạo con trỏ
    cursor = cnx.cursor(buffered=True)
    
    # Lấy danh sách các tên cột từ đối tượng con trỏ
    query = f"SELECT * FROM {table_name} LIMIT 1"
    cursor.execute(query)
    
    # Lấy danh sách các tên cột từ đối tượng con trỏ
    column_names = [column[0] for column in cursor.description]
    for index, row in df.iterrows():
        
        data = row.values.tolist()
        sql = query_sql(data,column_names,table_name)
        # print(sql)
        try:
            cursor.execute(sql)
            cnx.commit()
        except:
            pass
    cursor.close()
    cnx.close()

def updatedatabase_dubao(tab_name,matram,value):
    tg = '2024-01-25 00:00:00'
    # Tạo kết nối
    cnx = creat_cxn()
    # Tạo con trỏ
    cursor = cnx.cursor(buffered=True)
    sql = 'UPDATE {} SET SLDUNGDUOC = "{}",SLGOC = "{}" WHERE thoigian = "{}" and Matram = "{}"'.format(tab_name,value,value,tg,matram)
    print(sql)
    cursor.execute(sql)
    cnx.commit()
    cursor.close()
    cnx.close()
    
    
    
    
    
    
selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value


duyettin = None
def set_selected_duyet(value):
    global duyettin
    duyettin = value

def custom_round(value):
    if value != '-' and value % 1 == 0.5:
        return round(value + 0.1,0)
    return int(round(value,0))

def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def sobantin():
    now= datetime.now()
    bd = datetime((now-timedelta(days=365)).year,12,15)
    hs = now - bd
    sbt = int(hs.days*2)
    return sbt

def tin_tvhn():
    now = datetime.now()
    odoc = Document('TINMAU/ST_TVHN.docx')
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    # so ban tin
    sbt = sobantin()
    # print(sbt)
    if now.hour > 13:
        p=0
        tgpt = datetime(now.year,now.month,now.day,19,30)
    else:
        tgpt = datetime(now.year,now.month,now.day,7,30)
        p=1
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số: '+ str(sbt-p) + '/DBTVST2-ĐQNAM'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Nam' in dl:
                pr.text=''
                ntn = 'Quảng Nam, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
    

    for pr in odoc.paragraphs:
        dl = pr.text
        if 'Bảng 1: Lượng mưa (mm) từ 19 giờ ngày 20 đến 07 giờ ngày 21/3' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 1: Lượng mưa (mm) từ {} giờ ngày {} đến {} giờ ngày {}'.format((tgpt - timedelta(hours=12)).strftime('%H'),(tgpt - timedelta(hours=12)).strftime('%d/%m'),tgpt.strftime('%H'),tgpt.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.font.size = Pt(13)
        elif 'Bảng 2: Số liệu mực nước lúc 7h ngày 21/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 2: Số liệu mực nước lúc {} ngày {}'.format(tgpt.strftime('%H'),tgpt.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.font.size = Pt(13)
        elif 'Dự báo viên' in dl:
            # ban tin tiep theo
            ntn = 'Dự báo viên: {}'.format(selected_value)
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            # run.italic = True
            run.font.size = Pt(13)              
            
        elif 'Tin phát lúc' in dl:
            # ban tin tiep theo
            ntn = 'Tin phát lúc: {} giờ 30 phút.'.format(tgpt.strftime('%H'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)       

    # bang so 1
    odoc.tables[1].cell(0,1).paragraphs[0].add_run((tgpt - timedelta(hours=12)).strftime('%Hh') +' - '+(tgpt - timedelta(hours=6)).strftime('%Hh') ).font.size = Pt(13)        
    odoc.tables[1].cell(0,2).paragraphs[0].add_run((tgpt - timedelta(hours=6)).strftime('%Hh') +' - '+tgpt.strftime('%Hh') ).font.size = Pt(13)
    
    odoc.tables[5].cell(1,1).paragraphs[0].add_run((tgpt + timedelta(hours=12)).strftime('%Hh/%d/%m')).font.size = Pt(13)
    odoc.tables[5].cell(1,2).paragraphs[0].add_run((tgpt + timedelta(hours=24)).strftime('%Hh/%d/%m')).font.size = Pt(13)
    
    # lay so lieu mua
    pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/QNAM.accdb'
    FileName=(pth25)
    cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
    query = "SELECT * FROM mua"
    mua = pd.read_sql(query, cnxn)
    mua = mua[mua['thoigian'] > tgpt - timedelta(hours=12.5)]
    mua.set_index('thoigian',inplace= True)
    mua = mua.astype(float)
    muatong = mua.sum()
    nhanxetmua = ''
    if muatong.sum() ==0:
        nhanxetmua = 'Không mưa'
    else:
        kieumua = []
        for p_mua in muatong:
            if 5 <= p_mua and p_mua < 10:
                kieumua.append('mưa vừa') 
                break
        for p_mua in muatong:
            if 10 <= p_mua and p_mua < 20:
                kieumua.append('mưa to') 
                break    
        for p_mua in muatong:    
            if 20 <= p_mua :
                kieumua.append('mưa rất to') 
                break           
        if len(kieumua) == 1:
            if 'mưa vừa' in kieumua:
                nhanxetmua = 'Có mưa, có nơi mưa vừa.'
            elif 'mưa to' in kieumua:
                nhanxetmua = 'Có mưa, có nơi to.'
            elif 'mưa rất to' in kieumua:
                nhanxetmua = 'Có mưa, có nơi mưa rất to.'
        elif len(kieumua) == 2:
            if 'mưa vừa' in kieumua and 'mưa to' in kieumua:
                nhanxetmua = 'Có mưa, có nơi mưa vừa đến mưa to.'
            elif 'mưa to' in kieumua and 'mưa rất to' in kieumua :
                nhanxetmua = 'Có mưa, có nơi mưa to đến rất to.'
            elif 'mưa vừa' in kieumua and 'mưa rất to' in kieumua:
                nhanxetmua = 'Có mưa, có nơi mưa to đến rất to.'        
        elif len(kieumua) == 3:  
            nhanxetmua = 'Có mưa, mưa vừa, có nơi mưa to đến rất to.'      
        
    # print(len(muatong))
    
    
    mua6h = mua.rolling(6,min_periods=1).sum()
    mua6h = mua6h.loc[(mua6h.index.hour==1) | (mua6h.index.hour==7) | (mua6h.index.hour==13) | (mua6h.index.hour==19)]
    # print(mua6h)
    mua6h = mua6h.applymap("{0:.1f}".format)
    mua6h = mua6h.astype(str)
    mua6h =mua6h.replace('0.0','-')
    mua6h =mua6h.replace('nan','-')
    # print(mua6h)
    mua6h = mua6h.applymap(lambda x: custom_round(float(x)) if x != '-' else '-')
    # print(mua6h)
    mua6h = mua6h.astype(str)
    # mua6h =mua6h.replace('0.0','-')
    # print(mua6h)
    
    for pr in odoc.paragraphs:
        dl = pr.text
        if '1.1. Tình hình thời tiết: ' in dl:
            # ban tin tiep theo
            ntn = '1.1. Tình hình thời tiết: '
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)   
            run = pr.add_run(nhanxetmua)
            run.font.size = Pt(13) 
            break
            
    
    
    
# tracang  traleng  tranam2  trabui  tramai  tratap  dapsongtranh  tralinh  tragiac  tradon  travan
    odoc.tables[1].cell(1,1).paragraphs[0].add_run(mua6h['tralinh'].iloc[0]).font.size = Pt(13)
    try: 
        odoc.tables[1].cell(1,2).paragraphs[0].add_run(mua6h['tralinh'].iloc[1]).font.size = Pt(13) 
    except:
        pass
    odoc.tables[1].cell(2,1).paragraphs[0].add_run(mua6h['tranam2'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(2,2).paragraphs[0].add_run(mua6h['tranam2'].iloc[1]).font.size = Pt(13) 
    except:
        pass
    
    odoc.tables[1].cell(3,1).paragraphs[0].add_run(mua6h['travan'].iloc[0]).font.size = Pt(13)
    try: 
        odoc.tables[1].cell(3,2).paragraphs[0].add_run(mua6h['travan'].iloc[1]).font.size = Pt(13) 
    except:
        pass       
    odoc.tables[1].cell(4,1).paragraphs[0].add_run(mua6h['tracang'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(4,2).paragraphs[0].add_run(mua6h['tracang'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(5,1).paragraphs[0].add_run(mua6h['tramai'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(5,2).paragraphs[0].add_run(mua6h['tramai'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(6,1).paragraphs[0].add_run(mua6h['tragiac'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(6,2).paragraphs[0].add_run(mua6h['tragiac'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(7,1).paragraphs[0].add_run(mua6h['tradon'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(7,2).paragraphs[0].add_run(mua6h['tradon'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(8,1).paragraphs[0].add_run(mua6h['traleng'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(8,2).paragraphs[0].add_run(mua6h['traleng'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(9,1).paragraphs[0].add_run(mua6h['dapsongtranh'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(9,2).paragraphs[0].add_run(mua6h['dapsongtranh'].iloc[1]).font.size = Pt(13) 
    except:
        pass    
    odoc.tables[1].cell(10,1).paragraphs[0].add_run(mua6h['trabui'].iloc[0]).font.size = Pt(13)
    try:  
        odoc.tables[1].cell(10,2).paragraphs[0].add_run(mua6h['trabui'].iloc[1]).font.size = Pt(13) 
    except:
        pass
    # mua du bao vao bang
    mua_db = pd.read_excel('DATA/windy_db.xlsx')
    mua_db = mua_db[['time','muaTra Doc','muaTrung Luu','muaThuong Luu']]
    mua_db.set_index('time',inplace=True)
    mua_db = mua_db[(mua_db.index > tgpt - timedelta(minutes=30)) & (mua_db.index <= tgpt + timedelta(hours=23.5))]
    odoc.tables[3].cell(1,1).paragraphs[0].add_run("{0:.0f}".format(mua_db['muaTra Doc'].sum())).font.size = Pt(13)
    odoc.tables[3].cell(1,2).paragraphs[0].add_run("{0:.0f}".format(mua_db['muaTrung Luu'].sum())).font.size = Pt(13)
    odoc.tables[3].cell(1,3).paragraphs[0].add_run("{0:.0f}".format(mua_db['muaThuong Luu'].sum())).font.size = Pt(13)
    
    if duyettin=='Trương Tuyến':
        picture_filename = 'chuky/kydau_tuyen.png'

        
    elif duyettin=='Nguyễn Đình Huấn':
        picture_filename = 'chuky/kydau_huan.png'

    # chen chu ky
    table = odoc.tables[6]
            # Access the specific cell where you want to insert the picture
    cell = table.cell(0, 1)
    table.cell(0, 1).text=''
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.clear()

    # Add a picture to the cell
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # cell_paragraph.add_run(chucdanh + '\n').bold=True
    cell_paragraph.add_run().add_picture(picture_filename, width=Inches(2.0), height=Inches(1.5))
    # cell_paragraph.add_run("\n"+duyettin).bold=True
    
    pth = read_txt('path_tin/DRHN.txt') + '/QNAM_DBKTTV_STRANH_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    odoc.save(pth)
    # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
def tin_tv_load():
    now = datetime.now()
    if now.hour > 13:
        tgpt = datetime(now.year,now.month,now.day,19,30)
    else:
        tgpt = datetime(now.year,now.month,now.day,7,30)
    pth = read_txt('path_tin/DRHN.txt') + '/QNAM_DBKTTV_STRANH_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    if os.path.exists(pth):        
        odoc = Document(pth)
        style = odoc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
    
        pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/QNAM.accdb'
        FileName=(pth25)
        cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
        query = "SELECT thoigian,mucnuocthuongluu,qdenho FROM thuyvan"
        mucnuoc = pd.read_sql(query, cnxn)
        mucnuoc['thoigian'] = pd.to_datetime(mucnuoc['thoigian'])
        h_ho = mucnuoc[mucnuoc['thoigian']==(tgpt - timedelta(minutes=30))]
        mucnuoc = h_ho['mucnuocthuongluu'].values[0]
        qve = h_ho['qdenho'].values[0]
        
        df = TTB_API_mucnuoc()
        df = df.interpolate(method='linear')
        h_giaothuy = df[df.index ==(tgpt - timedelta(minutes=30))]['Giao Thuy'].values[0]


        odoc.tables[2].cell(1,1).text = ''
        odoc.tables[2].cell(2,1).text = ''
        odoc.tables[2].cell(1,2).text = ''  
        odoc.tables[2].cell(2,2).text = ''
        
        odoc.tables[2].cell(1,1).paragraphs[0].add_run(mucnuoc)
        odoc.tables[2].cell(2,1).paragraphs[0].add_run(qve)
        odoc.tables[2].cell(1,2).paragraphs[0].add_run(str(h_giaothuy)) 
        
        for i in range(1,3):
            for j in range(1,3):
                pr = odoc.tables[2].cell(i,j).paragraphs[0]
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        odoc.save(pth)
        # convert(pth,pth.replace('.docx','.pdf'))
        messagebox.showinfo('Thông báo','OK!')
    else:
        messagebox.showinfo('Thông báo','Chưa tạo tin nền KT!')
    
def nghiemthu_tvhn():
    now = datetime.now()
    if now.hour > 13:
        tgpt = datetime(now.year,now.month,now.day,19,30)
    else:
        tgpt = datetime(now.year,now.month,now.day,7,30)
    pth = read_txt('path_tin/DRHN.txt') + '/QNAM_DBKTTV_STRANH_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    if os.path.exists(pth):        
        odoc = Document(pth)
        style = odoc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        
        if duyettin=='Trương Tuyến':
            picture_filename = 'chuky/tuyen.png'
            chucdanh = 'GIÁM ĐỐC'
            
        elif duyettin=='Nguyễn Đình Huấn':
            picture_filename = 'chuky/huan.png'
            chucdanh = 'KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC'
        
        # chen chu ky
        table = odoc.tables[6]
        cell = table.cell(0, 1)
        table.cell(0, 1).text=''
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.clear()

        # Add a picture to the cell
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_paragraph.add_run(chucdanh + '\n').bold=True
        cell_paragraph.add_run().add_picture(picture_filename, width=Inches(1.3), height=Inches(1.0))
        cell_paragraph.add_run("\n"+duyettin).bold=True
        odoc.save(pth)
        uploaddb_tvhn()
        messagebox.showinfo('Thông báo','OK!')
    else:
        messagebox.showinfo('Thông báo','Chưa tạo tin nền KT!')
        
def uploaddb_tvhn():
    pth = tim_file(read_txt('path_tin/DRHN.txt'),'.docx')
    if os.path.exists(pth):        
        odoc = Document(pth)
        # mua du bao
        muabd24_dap = odoc.tables[3].cell(1,1).text
        # muabd24_trungluu = odoc.tables[3].cell(1,1).text
        # muabd24_thuongluu = odoc.tables[3].cell(1,1).text
        # qbd du bao
        qbd =  odoc.tables[4].cell(1,1).text
        qbd = qbd.split('-')
        trung_binh = np.mean([int(x) for x in qbd])
        print(trung_binh)
        # hdb du bao
        hdb = odoc.tables[5].cell(2,2).text
        updatedatabase_dubao('ho_dakdrinh_mucnuoc','muadb24',str(muabd24_dap))
        updatedatabase_dubao('ho_dakdrinh_mucnuoc','qdubao',str(trung_binh))
        updatedatabase_dubao('ho_dakdrinh_mucnuoc','hdb',str(hdb))
        
        
        
        

        
        

    