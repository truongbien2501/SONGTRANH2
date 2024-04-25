from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
from func.load_data import mo_word
import pyodbc
from func.CDH_TTB_API import TTB_API_mucnuoc10day
selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value

duyettin = None

def set_selected_duyet(value):
    global duyettin
    duyettin = value

def custom_round(value):
    if value != '-' and value % 1 == 0.5:
        return round(value + 0.1,0)
    return int(round(value,0))

def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def tin_nenKT_thang():
    now = datetime.now()
    odoc = Document('TINMAU/ST_TVHT.docx')
    # odoc = Document(r'D:\PM_PYTHON\SONGTRANH\TINMAU\ST_TVHV10.docx')
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    # so ban tin
    # sbt = 25
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số: '+ str(now.strftime('%m')) + '/DBHDST2-ĐQNAM'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Nam' in dl:
                pr.text=''
                ntn = 'Quảng Nam, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
            

    
    for pr in odoc.paragraphs:
        dl = pr.text
        if 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN DÀI' in dl:
            # ban tin tiep theo
            ntn = 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN DÀI'
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(14)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '(Tháng 4 năm 2024)' in dl:
            # ban tin tiep theo
            ntn = '(Tháng {} năm 2024)'.format(now.strftime('%m'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '1. Dự báo thời tiết tháng 4/2024 trên lưu vực Sông Tranh' in dl:
            # ban tin tiep theo
            ntn = '1. Dự báo thời tiết tháng 4/2024 trên lưu vực Sông Tranh'.format(now.strftime('%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif 'Tháng 4 năm 2024, lưu vực Sông Tranh chịu ảnh hưởng của hệ thống thời tiết sau' in dl:
            # ban tin tiep theo
            ntn = 'Tháng {} năm {}, lưu vực Sông Tranh chịu ảnh hưởng của hệ thống thời tiết sau: phần phía Nam của rãnh áp thấp có trục qua khu vực Bắc Trung Bộ, vùng áp thấp phía Tây phát triển mở rộng về phía Đông Nam, đới gió mùa Tây Nam hoạt động mạnh dần, rãnh áp thấp có trục qua khu vực Nam Trung Bộ nâng trục dần lên qua khu vực, hoạt động của áp cao cận nhiệt đới trên cao.'.format(now.strftime('%m'),now.strftime('%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif 'Xu thế lượng mưa tháng 4 năm 2024 ' in dl:
            # ban tin tiep theo
            ntn = 'Xu thế lượng mưa tháng {} năm {} tại lưu vực Sông Tranh phổ biến ở mức xấp xỉ hoặc thấp hơn giá trị TBNN từ 10 - 15%.'.format(now.strftime('%m'),now.strftime('%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT                        
        elif 'Bảng 1: Tổng lượng mưa tháng 3/2024 và Dự báo lượng mưa tháng 4/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 1: Tổng lượng mưa tháng {} và Dự báo lượng mưa tháng {}'.format((now-timedelta(days=29)).strftime('%m/%Y'),now.strftime('%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold=False
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Tháng 04 năm 2024, dòng chảy đến hồ Thủy điện Sông Tranh 2' in dl:
            ntn = 'Tháng {} năm {}, dòng chảy đến hồ Thủy điện Sông Tranh 2 có biến đổi chậm.'.format(now.strftime('%m'),now.strftime('%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT    
        elif 'Bảng 2: Đặc trưng lưu lượng tháng 3/2024 và dự báo lưu lượng tháng 4/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 2: Đặc trưng lưu lượng tháng {} và dự báo lưu lượng tháng {}'.format((now-timedelta(days=29)).strftime('%m/%Y'),now.strftime('%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            run.bold=False
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        elif 'Thời gian ban hành bản tin tiếp theo' in dl:
            thangtieptheo = now + timedelta(days=31)
            # ban tin tiep theo
            ntn = 'Thời gian ban hành bản tin tiếp theo: 16 giờ 15 phút ngày {}.'.format(datetime(thangtieptheo.year,thangtieptheo.month,1).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)              
        elif 'Dự báo viên' in dl:
            # ban tin tiep theo
            ntn = 'Dự báo viên: {}'.format(selected_value)
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)  
        elif 'Phụ lục: Tổng lượng mưa (mm) tháng 3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Phụ lục: Tổng lượng mưa (mm) tháng {}'.format((now-timedelta(days=29)).strftime('%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    thangtruoc = now - timedelta(days=28)        
    bd_mua = datetime(thangtruoc.year,thangtruoc.month,1,20) - timedelta(days=1)
    # lay so lieu mua
    pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/QNAM.accdb'
    # pth25 = r'D:\PM_PYTHON\SONGTRANH\DATA\QNAM.accdb'
    FileName=(pth25)
    cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
    query = "SELECT * FROM mua"
    mua = pd.read_sql(query, cnxn)
    mua = mua[(mua['thoigian'] >=bd_mua) & (mua['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,19))]
    mua.set_index('thoigian',inplace=True)
    mua = mua.astype(float)
    muathang = mua.sum()
    muathang = muathang.replace(0.0,'-')
    print(muathang)
    muathang = muathang.map(lambda x: custom_round(float(x)) if x != '-' else '-')
    
    
    dt_mua = muathang.replace('-',0)
    max1 = max( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    min1 = min( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    max2 = max( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    min2 = min( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    max3 = max( dt_mua['trabui'], dt_mua['dapsongtranh'])
    min3 = min( dt_mua['trabui'], dt_mua['dapsongtranh'])
    

    
    # print(muangay)

        
    muathang = muathang.astype(str)
    # print(mua10)
    # bang top hop mua so 1   
    odoc.tables[1].cell(1,0).text= 'Tổng lượng mưa tháng {}'.format(thangtruoc.strftime('%m/%Y'))
    odoc.tables[1].cell(1,1).text= str(min1) + ' - ' + str(max1)
    odoc.tables[1].cell(1,2).text= str(min2) + ' - ' + str(max2)
    odoc.tables[1].cell(1,3).text= str(min3) + ' - ' + str(max3)   
    odoc.tables[1].cell(2,0).text= 'Tổng lượng mưa tháng {}'.format(now.strftime('%m/%Y'))
    # dieu chinh lai format cho bang
    for a in range(1,3):
        for b in range(0,4):
            try:
                for pr in odoc.tables[1].cell(a,b).paragraphs:
                    pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in pr.runs:
                        run.font.size = Pt(12)
                        run.bold = False  
            except:
                pass




    odoc.tables[4].cell(1,2).text= muathang['tralinh']
    odoc.tables[4].cell(2,2).text= muathang['tranam2']
    odoc.tables[4].cell(3,2).text= muathang['travan']
    odoc.tables[4].cell(4,2).text= muathang['tracang']
    odoc.tables[4].cell(5,2).text= muathang['tramai']
    odoc.tables[4].cell(6,2).text= muathang['tragiac']
    odoc.tables[4].cell(7,2).text= muathang['tradon']
    odoc.tables[4].cell(8,2).text= muathang['traleng']
    odoc.tables[4].cell(9,2).text= muathang['dapsongtranh']
    odoc.tables[4].cell(10,2).text= muathang['trabui']
    # dieu chinh lai format cho bang
    for a in range(1,11):
        for b in range(2,3):
            try:
                for pr in odoc.tables[4].cell(a,b).paragraphs:
                    pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in pr.runs:
                        run.font.size = Pt(12)
                        run.bold = False  
            except:
                pass


    # bang luu luong nuoc
    odoc.tables[2].cell(0,1).text= 'Tháng {}'.format(thangtruoc.strftime('%m/%Y'))
    odoc.tables[2].cell(0,2).text= 'Tháng {}'.format(now.strftime('%m/%Y'))
    # dieu chinh lai format cho bang
    for a in range(0,1):
        for b in range(1,3):
            try:
                for pr in odoc.tables[2].cell(a,b).paragraphs:
                    pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in pr.runs:
                        run.font.size = Pt(12)
                        run.bold = False  
            except:
                pass
    try:
        mucnuoc = TTB_API_mucnuoc10day()
        mucnuoc = mucnuoc.reset_index(False)
        mucnuoc.rename(columns={'time':'thoigian','qden':'qdenho'},inplace=True)
        cnxn.close()
    except:
        query = "SELECT thoigian,qdenho FROM thuyvan"
        mucnuoc = pd.read_sql(query, cnxn)
        cnxn.close()
    # print(bd_mua + timedelta(hours=4))
    mucnuoc = mucnuoc[(mucnuoc['thoigian'] >=(bd_mua + timedelta(hours=4))) & (mucnuoc['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,23))]

    mucnuoc.set_index('thoigian',inplace=True)
    # mucnuoc.to_excel('kiemtrah.xlsx')
    mucnuoc = mucnuoc.astype(float)
    # bang thuc do luu luong 2
    odoc.tables[2].cell(1,1).text= '{0:.1f}'.format(mucnuoc['qdenho'].mean())
    odoc.tables[2].cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # # print(mua6h)
    pth = read_txt('path_tin/DRHD.txt') + '/BT1T_ST2_QNAM_{}.docx'.format(now.strftime('%Y%m%d')) 
    odoc.save(pth)
    # # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    