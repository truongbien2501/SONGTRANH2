from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
from func.load_data import mo_word
from func import load_data
from win32com import client
import pyodbc
from func.CDH_TTB_API import TTB_API_mucnuoc_lu
selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value

def sosanh_BD(h,tentram):
    if tentram == 'Nong Son':
        bd1,bd2,bd3=11,13,15
    elif tentram == 'Cau Lau':
        bd1,bd2,bd3=2,3,4
    elif tentram == 'Ai Nghia':
        bd1,bd2,bd3=6.5,8,9
    elif tentram == 'Giao Thuy':
        bd1,bd2,bd3=6.5,7.5,8.8

    if h > bd3:
        nx = "Trên BĐ3: " + str(h - bd3) + 'm'
    elif h == bd3:
        nx = "Ở mức BĐ3"
    elif bd2 < h < bd3 and (bd3-h) < (h-bd2):
        nx = "Dưới BĐ3: " + str(bd3-h)+ 'm'
    elif bd2 < h < bd3 and (bd3-h) >= (h-bd2):
        nx = "Trên BĐ2: " + str(h-bd2)+ 'm'
    elif h == bd2:
        nx = "Ở mức BĐ2"
    elif bd1 < h < bd2 and (bd2-h) < (h-bd1):
        nx = "Dưới BĐ2: " + str(bd2-h)+ 'm'
    elif bd1 < h < bd2 and (bd2-h) >= (h-bd1):
        nx = "Trên BĐ1: " + str(h-bd1)+ 'm'
    elif h == bd1:
        nx = "Ở mức BĐ1"
    elif h < bd1:
        nx = "Dưới BĐ1: " +  str(bd1-h)+ 'm'
    return nx

def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def thoigianphattin():
    now = datetime.now()
    if now >= datetime(now.year,now.month,now.day,0) and  now <= datetime(now.year,now.month,now.day,2,30):
        tgpt = datetime(now.year,now.month,now.day,1,30)
    elif now >= datetime(now.year,now.month,now.day,3) and  now <= datetime(now.year,now.month,now.day,5,30):
        tgpt = datetime(now.year,now.month,now.day,4,30)    
    elif now >= datetime(now.year,now.month,now.day,6) and  now <= datetime(now.year,now.month,now.day,9,0):
        tgpt = datetime(now.year,now.month,now.day,7,30)    
    elif now >= datetime(now.year,now.month,now.day,9) and  now <= datetime(now.year,now.month,now.day,12):
        tgpt = datetime(now.year,now.month,now.day,10,30)        
    elif now >= datetime(now.year,now.month,now.day,12) and  now <= datetime(now.year,now.month,now.day,14,30):
        tgpt = datetime(now.year,now.month,now.day,13,30) 
    elif now >= datetime(now.year,now.month,now.day,15) and  now <= datetime(now.year,now.month,now.day,17,30):
        tgpt = datetime(now.year,now.month,now.day,16,30)     
    elif now >= datetime(now.year,now.month,now.day,18) and  now <= datetime(now.year,now.month,now.day,20,30):
        tgpt = datetime(now.year,now.month,now.day,19,30)  
    elif now >= datetime(now.year,now.month,now.day,21) and  now <= datetime(now.year,now.month,now.day,23,30):
        tgpt = datetime(now.year,now.month,now.day,22,30)  
    return tgpt

def sobt():
    pth = tim_file(read_txt('path_tin/LULU.txt'),'.docx')
    # print(pth)
    ttpt = thoigianphattin()
    if ttpt.strftime('%Y%m%d_%H30') in pth:
        os.remove(pth)
        messagebox.showinfo('Thong bao','Đã xóa file tồn tại' + pth.split('\\')[-1])
        pth = tim_file(read_txt('path_tin/LULU.txt'),'.docx')
    odoc = Document(pth)
    for a in odoc.tables[0].cell(0,0).paragraphs:
        if 'Số' in a.text:
            dl = str(a.text)
            sbt = dl[dl.index(':')+1:dl.index('/')]
    return int(sbt) + 1

def xacdinhngaydb():
    now = datetime.now()
    for a in range(3,10):
        tttt = now + timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def xacdinhngaydaqua():
    now = datetime.now()
    for a in range(3,10):
        tttt = now - timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def custom_round(value):
    if value != '-' and value % 1 == 0.5:
        return round(value + 0.1,0)
    return int(round(value,0))


def tin_nenKT_lulu():
    tgpt = thoigianphattin()
    sbt = sobt()
    if sbt < 10:
        sbt = '0' + str(sbt)
    else:
        sbt = str(sbt)
    now = datetime.now()
    pth = tim_file(read_txt('path_tin/LULU.txt'),'.docx')
    odoc = Document(pth)
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số: '+ str(sbt) + '/TLST2-ĐKTTVQN'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Nam' in dl:
                pr.text=''
                ntn = 'Quảng Nam, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
        
    # lay so lieu mua
    pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/QNAM.accdb'
    # pth25 = r'D:\PM_PYTHON\SONGTRANH\DATA\QNAM.accdb'
    FileName=(pth25)
    cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
    query = "SELECT * FROM mua"
    mua = pd.read_sql(query, cnxn)
    
    mua = mua[(mua['thoigian'] >=(tgpt-timedelta(hours=3))) & (mua['thoigian'] <= (tgpt-timedelta(hours=0.5)))]
    mua.set_index('thoigian',inplace=True)
    mua = mua.astype(float)
    mua3h = mua.sum()
    nhanxetmua = ''
    if mua3h.sum() == 0:
        nhanxetmua = 'không mưa'
    else:
        kieumua = []
        for p_mua in mua3h:
            if 0 < p_mua and p_mua < 5:
                kieumua.append('mưa nhỏ') 
                break
        for p_mua in mua3h:
            if 5 <= p_mua and p_mua < 10:
                kieumua.append('mưa vừa') 
                break
        for p_mua in mua3h:
            if 10 <= p_mua and p_mua < 20:
                kieumua.append('mưa to') 
                break    
        for p_mua in mua3h:    
            if 20 <= p_mua :
                kieumua.append('mưa rất to') 
                break           
        if len(kieumua) == 1:
            if 'mưa nhỏ' in kieumua:
                nhanxetmua = 'có mưa nhỏ.'
            elif 'mưa vừa' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa vừa.'
            elif 'mưa to' in kieumua:
                nhanxetmua = 'có mưa, có nơi to.'
            elif 'mưa rất to' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa rất to.'
        elif len(kieumua) == 2:
            if 'mưa nhỏ' in kieumua and 'mưa vừa' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa vừa.'
            elif 'mưa nhỏ' in kieumua and 'mưa to' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa to.'
            elif 'mưa nhỏ' in kieumua and 'mưa rất to' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa rất to.'                                        
            elif 'mưa vừa' in kieumua and 'mưa to' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa vừa đến mưa to.'
            elif 'mưa vừa' in kieumua and 'mưa rất to' in kieumua:
                nhanxetmua = 'có mưa, có nơi mưa to đến rất to.'
            elif 'mưa to' in kieumua and 'mưa rất to' in kieumua :
                nhanxetmua = 'có mưa, có nơi mưa to đến rất to.'        
        elif len(kieumua) == 3:  
            nhanxetmua = 'có mưa, mưa vừa, có nơi mưa to đến rất to.'     
    
    mua3h = mua3h.replace(0,'-')
    mua3h = mua3h.apply(lambda x: custom_round(float(x)) if x != '-' else '-')
    mua3h = mua3h.astype(str)
    # print(mua3h)        
        
    odoc.tables[1].cell(1,1).text = mua3h['tralinh']
    odoc.tables[1].cell(1,2).text = mua3h['tranam2']
    odoc.tables[1].cell(1,3).text = mua3h['travan']
    odoc.tables[1].cell(1,4).text = mua3h['tracang']
    odoc.tables[1].cell(1,5).text = mua3h['tramai']
    odoc.tables[1].cell(1,6).text = mua3h['tragiac']
    odoc.tables[1].cell(1,7).text = mua3h['tradon']
    odoc.tables[1].cell(1,8).text = mua3h['traleng']
    odoc.tables[1].cell(1,9).text = mua3h['dapsongtranh']
    odoc.tables[1].cell(1,10).text = mua3h['trabui']
    
    
    # load so lieu muc nuoc 
    dfh = TTB_API_mucnuoc_lu()
    dfh = dfh.interpolate(method='linear')
    h_ns = dfh[dfh.index ==(tgpt - timedelta(minutes=30))]['Nong Son'].values[0]
    h_cl = dfh[dfh.index ==(tgpt - timedelta(minutes=30))]['Cau Lau'].values[0]
    h_st = dfh[dfh.index ==(tgpt - timedelta(minutes=30))]['mucnuoc'].values[0]
    q_den = dfh[dfh.index ==(tgpt - timedelta(minutes=30))]['qden'].values[0]
    for pr in odoc.paragraphs:
        dl = pr.text
        if 'TIN LŨ VỀ HỒ THUỶ ĐIỆN SÔNG TRANH' in dl:
            # ban tin tiep theo
            ntn = 'TIN LŨ VỀ HỒ THUỶ ĐIỆN SÔNG TRANH 2'
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(14)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '1. Tình hình mưa (đơn vị mm):' in dl:
            # ban tin tiep theo
            ntn = '1. Tình hình mưa (đơn vị mm): '
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            run.bold =True
            run.italic =True
            # pr.alignment = WD_PARAGRAPH_ALIGNMENT.Left
            ntn = 'Trong 3 giờ qua ({} - {}) lưu vực sông tranh {}'.format((tgpt-timedelta(hours=3.5)).strftime('%Hh %d/%m'),tgpt.strftime('%Hh %d/%m/%Y'),nhanxetmua)
            run = pr.add_run(ntn)
            
        elif 'Bảng 1' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 1: ' 
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.italic =True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ntn = 'Lượng mưa từ {} - {}'.format((tgpt-timedelta(hours=3.5)).strftime('%Hh ngày %d/%m'),(tgpt-timedelta(hours=0.5)).strftime('%Hh ngày %d/%m/%Y'))
            run = pr.add_run(ntn)
        elif 'Bảng 2' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 2: '
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.italic =True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
            ntn = 'Mực nước lúc {} ngày {}'.format((tgpt-timedelta(hours=0.5)).strftime('%Hh'),(tgpt-timedelta(hours=0.5)).strftime('%d/%m/%Y'))
            run = pr.add_run(ntn)
        elif 'Hiện tại lưu lượng nước về hồ Thủy điện Sông Tranh 2' in dl:
            # ban tin tiep theo
            ntn = 'Hiện tại lưu lượng nước về hồ Thủy điện Sông Tranh 2 dao động ở mức cao. Lúc {} ngày {}, lưu lượng nước về hồ {}m3/s, mực nước hồ {}m.'.format(tgpt.strftime('%Hh'),tgpt.strftime('%d/%m/%Y'),'{0:.2f}'.format(h_st),'{0:.2f}'.format(q_den))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.italic =True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
            ntn = 'Dự báo mưa thời đoạn 6 giờ từ {} đến {}'.format(tgpt.strftime('%Hh ngày %d/%m'),(tgpt + timedelta(hours=23.5)).strftime('%Hh ngày %d/%m/%Y'))
            run = pr.add_run(ntn)            
        elif 'Bảng 3' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 3: '
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.italic =True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
            ntn = 'Dự báo mưa thời đoạn 6 giờ từ {} đến {}'.format(tgpt.strftime('%Hh ngày %d/%m'),(tgpt + timedelta(hours=23.5)).strftime('%Hh ngày %d/%m/%Y'))
            run = pr.add_run(ntn)
        elif 'Bảng 4' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 4: '
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
            ntn = 'Dự báo lưu lượng đến hồ thời đoạn 6 giờ từ {} - {}'.format(tgpt.strftime('%Hh ngày %d/%m'),(tgpt + timedelta(hours=23.5)).strftime('%Hh ngày %d/%m'))
            run = pr.add_run(ntn)
        elif 'Tin phát lúc' in dl:
            # ban tin tiep theo
            ntn = 'Tin phát lúc: {} giờ {} phút.'.format(tgpt.strftime('%H'),tgpt.strftime('%M'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.italic = True
            run.font.size = Pt(13)            
        # elif 'Dự báo viên' in dl:
        #     # ban tin tiep theo
        #     ntn = 'Dự báo viên: {}'.format(selected_value)
        #     pr.text  =''
        #     run = pr.add_run(ntn)
        #     # run.bold = True
        #     run.italic = True
        #     run.font.size = Pt(13)          

    # print(dfh)
    # bang thuc do muc nuoc cau lau va nong son
    odoc.tables[2].cell(1,1).text = str(h_ns)
    odoc.tables[2].cell(1,2).text = str(h_cl)
    odoc.tables[2].cell(2,1).text = sosanh_BD(h_ns,'Nong Son')
    odoc.tables[2].cell(2,2).text = sosanh_BD(h_cl,'Cau Lau')            
            
            
    # bang 3 bang du bao mua thoi doan 6
    odoc.tables[3].cell(1,1).text = tgpt.strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5)).strftime('%Hh/%d')
    odoc.tables[3].cell(1,2).text = (tgpt + timedelta(hours=5.5)).strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5+6)).strftime('%Hh/%d')
    odoc.tables[3].cell(1,3).text = (tgpt + timedelta(hours=5.5+6)).strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5+12)).strftime('%Hh/%d')   
    odoc.tables[3].cell(1,4).text = (tgpt + timedelta(hours=5.5+12)).strftime('%Hh/%d/') + ' - ' + (tgpt + timedelta(hours=5.5+18)).strftime('%Hh/%d')
    odoc.tables[3].cell(2,1).text = odoc.tables[3].cell(2,2).text
    odoc.tables[3].cell(2,2).text = odoc.tables[3].cell(2,3).text
    odoc.tables[3].cell(2,3).text = odoc.tables[3].cell(2,4).text  
    odoc.tables[3].cell(2,4).text = ''  
    
    odoc.tables[4].cell(1,1).text = tgpt.strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5)).strftime('%Hh/%d')
    odoc.tables[4].cell(1,2).text = (tgpt + timedelta(hours=5.5)).strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5+6)).strftime('%Hh/%d')
    odoc.tables[4].cell(1,3).text = (tgpt + timedelta(hours=5.5+6)).strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5+12)).strftime('%Hh/%d')   
    odoc.tables[4].cell(1,4).text = (tgpt + timedelta(hours=5.5+12)).strftime('%Hh/%d') + ' - ' + (tgpt + timedelta(hours=5.5+18)).strftime('%Hh/%d')
    odoc.tables[4].cell(2,1).text = odoc.tables[4].cell(2,2).text
    odoc.tables[4].cell(2,2).text = odoc.tables[4].cell(2,3).text
    odoc.tables[4].cell(2,3).text = odoc.tables[4].cell(2,4).text      
    odoc.tables[4].cell(2,4).text = ''
            
    # dieu chinh lai format cho bang thong ke
    for t in range (1,5):
        for a in range(0,12):
            for b in range(0,12):
                try:
                    for pr in odoc.tables[t].cell(a,b).paragraphs:
                        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in pr.runs:
                            run.font.size = Pt(12)
                            run.bold = False  
                except:
                    pass
 
    pth = read_txt('path_tin/LULU.txt') + '/QNAM_TINLU_ST2_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    odoc.save(pth)
    messagebox.showinfo('Thông báo','OK!')
    
def tin_lulu_load():
    now = datetime.now()
    tgpt = thoigianphattin()
    pth_docx = read_txt('path_tin/LULU.txt') + '/DHC_LULU_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    odoc = Document(pth_docx)
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
   
    df = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',engine="openpyxl",sheet_name='DRHN')

    kt = now - timedelta(hours=24)
    df = df.iloc[1:,:21]
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    df['deta_t'] = df['time'].diff().dt.components['hours']
    # print(tgpt)
    if tgpt.hour==4 or tgpt.hour==10 or tgpt.hour==16 or tgpt.hour==22:
        df = df.loc[(df['time'] > (tgpt - timedelta(hours=18))) & (df['time'] <= (tgpt + timedelta(hours=26.5))) ]
        df.loc[df['time'] == (tgpt + timedelta(hours=2.5)), 'time'] = tgpt - timedelta(hours=0.5)
        # update lai khoang thoi gian
        df.loc[df['time'] == (tgpt - timedelta(hours=0.5)), 'deta_t'] = 3
        df.loc[df['time'] == (tgpt + timedelta(hours=8.5)), 'deta_t'] = 3
        
        # df_td = df.loc[(df['time'] > (tgpt - timedelta(hours=18))) & (df['time'] <= (tgpt + timedelta(hours=2.5))) ] # so lieu thuc do
        # df_db = df.loc[(df['time'] > tgpt) & (df['time'] <= (tgpt + timedelta(hours=24.5))) ] # so lieu du bao
    else:
        # df_td = df.loc[(df['time'] > (tgpt - timedelta(hours=24))) & (df['time'] <= tgpt) ] # so lieu thuc do
        # df_db = df.loc[(df['time'] > (tgpt + timedelta(hours=0.7+1+1))) & (df['time'] <= (tgpt + timedelta(hours=24))) ] # so lieu du bao
        df = df.loc[(df['time'] > (tgpt - timedelta(hours=24))) & (df['time'] <= (tgpt + timedelta(hours=24.5))) ]
    # print(df_td)
    # print(df_db)
    

    # df = df.loc[(df['time'] > (tgpt - timedelta(hours=18))) & (df['time'] <= (tgpt + timedelta(hours=24.5))) ]
    df['wden'] = (df['qdb'] - df['qdtdk'] - df['qcmdk'])*df['deta_t']*36/10**4
    
    # print(df["wtd"])
    # df["wtd"] = df["wden"].cumsum()
    # so lieu du bao
    df_db = df.loc[df['time'] >= (tgpt-timedelta(hours=0.5))] # so lieu du bao
    df_db["wtd"].iloc[1:] = df_db["wden"].iloc[1:].cumsum() + df_db["wtd"].iloc[0]
    # df_db["wtd"] =  df_db["wtd"].applymap("{0:.2f}".format)
    df_db["wtd"] =  df_db["wtd"].apply(lambda x: round(x, 2) if x is not None else x)
    df_h_w = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',sheet_name='Z-F-W')
    df_h_w = df_h_w[['H','W']]
    df_h_w.rename(columns={'W':'wtd'},inplace=True)
    df_h_w =df_h_w.iloc[3:,:]
    df_h_w["wtd"] =  df_h_w["wtd"].apply(lambda x: round(x, 2) if x is not None else x)
    # # print(df_h_w)
    # df_db =df_db.merge(df_h_w,how='left',on='wtd')
    # print(df["wtd"])
    df_db = df_db[1:]
    
    # so lieu tra khuc
    # print(df_db) 
    id =load_data.vitridat_hn_lulu() + 1
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        df_db[['wden','wtd']].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=9, header=False, index=False)
    
    
    df = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',engine="openpyxl",sheet_name='DRHN')

    kt = now - timedelta(hours=24)
    df = df.iloc[1:,:21]
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    df['deta_t'] = df['time'].diff().dt.components['hours']
    print(tgpt)
    if tgpt.hour==4 or tgpt.hour==10 or tgpt.hour==16 or tgpt.hour==22:
        df = df.loc[(df['time'] > (tgpt - timedelta(hours=18))) & (df['time'] <= (tgpt + timedelta(hours=26.5))) ]
        df.loc[df['time'] == (tgpt + timedelta(hours=2.5)), 'time'] = tgpt - timedelta(hours=0.5)
        # update lai khoang thoi gian
        df.loc[df['time'] == (tgpt - timedelta(hours=0.5)), 'deta_t'] = 3
        df.loc[df['time'] == (tgpt + timedelta(hours=8.5)), 'deta_t'] = 3
        
        # df_td = df.loc[(df['time'] > (tgpt - timedelta(hours=18))) & (df['time'] <= (tgpt + timedelta(hours=2.5))) ] # so lieu thuc do
        # df_db = df.loc[(df['time'] > tgpt) & (df['time'] <= (tgpt + timedelta(hours=24.5))) ] # so lieu du bao
    else:
        # df_td = df.loc[(df['time'] > (tgpt - timedelta(hours=24))) & (df['time'] <= tgpt) ] # so lieu thuc do
        # df_db = df.loc[(df['time'] > (tgpt + timedelta(hours=0.7+1+1))) & (df['time'] <= (tgpt + timedelta(hours=24))) ] # so lieu du bao
        df = df.loc[(df['time'] > (tgpt - timedelta(hours=24))) & (df['time'] <= (tgpt + timedelta(hours=24.5))) ]
    # print(df)
    df['Htd'] = df['Htd'].apply('{0:.2f}'.format)
    df['Htdtk'] = df['Htdtk'].apply('{0:.2f}'.format)
    df['Hmax'] = df['Hmax'].apply('{0:.2f}'.format)
    # df['Hdb'] = df['Hdb'].apply(lambda x: f"{x}")
    df['Hdb'] = df['Hdb'].apply('{0:.2f}'.format)
    # print(df)
    # df_qh = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',sheet_name='Z-F-W')
    # df_qh = df_qh[['H','W']]
    # df_qh.rename(columns={'W':'wtd'},inplace=True)
    # df_qh['wtd'] = df_qh['wtd'].apply('{0:.1f}'.format)

    # df = df.merge(df_qh,how='left',on='wtd')
    # df['H'] = df['H'].apply('{0:.2f}'.format)
    
    print(df)
    
    for t in range(3,7,3):
        for row in odoc.tables[t].rows[1:]:
            for cell in row.cells[1:]:
                cell.text = ''
    for row in odoc.tables[4].rows[1:]:
        for cell in row.cells[1:]:
            cell.text = ''
                # for paragraph in cell.paragraphs:
                #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    
    
    odoc.tables[3].cell(1,1).paragraphs[0].add_run(lamtron_Q(df['qtd'].iloc[0])).font.size = Pt(13)
    odoc.tables[3].cell(2,1).paragraphs[0].add_run(lamtron_Q(df['qtd'].iloc[1])).font.size = Pt(13) 
    odoc.tables[3].cell(3,1).paragraphs[0].add_run(lamtron_Q(df['qtd'].iloc[2])).font.size = Pt(13) 
    odoc.tables[3].cell(4,1).paragraphs[0].add_run(lamtron_Q(df['qtd'].iloc[3])).font.size = Pt(13) 
    
    odoc.tables[3].cell(1,2).paragraphs[0].add_run(lamtron_Q(df['qdt'].iloc[0])).font.size = Pt(13)
    odoc.tables[3].cell(2,2).paragraphs[0].add_run(lamtron_Q(df['qdt'].iloc[1])).font.size = Pt(13) 
    odoc.tables[3].cell(3,2).paragraphs[0].add_run(lamtron_Q(df['qdt'].iloc[2])).font.size = Pt(13) 
    odoc.tables[3].cell(4,2).paragraphs[0].add_run(lamtron_Q(df['qdt'].iloc[3])).font.size = Pt(13) 
    
    odoc.tables[3].cell(1,3).paragraphs[0].add_run(lamtron_Q(df['qxa'].iloc[0])).font.size = Pt(13)
    odoc.tables[3].cell(2,3).paragraphs[0].add_run(lamtron_Q(df['qxa'].iloc[1])).font.size = Pt(13) 
    odoc.tables[3].cell(3,3).paragraphs[0].add_run(lamtron_Q(df['qxa'].iloc[2])).font.size = Pt(13) 
    odoc.tables[3].cell(4,3).paragraphs[0].add_run(lamtron_Q(df['qxa'].iloc[3])).font.size = Pt(13)     
    
    odoc.tables[3].cell(1,4).paragraphs[0].add_run(lamtron_Q(df['qcm'].iloc[0])).font.size = Pt(13)
    odoc.tables[3].cell(2,4).paragraphs[0].add_run(lamtron_Q(df['qcm'].iloc[1])).font.size = Pt(13) 
    odoc.tables[3].cell(3,4).paragraphs[0].add_run(lamtron_Q(df['qcm'].iloc[2])).font.size = Pt(13) 
    odoc.tables[3].cell(4,4).paragraphs[0].add_run(lamtron_Q(df['qcm'].iloc[3])).font.size = Pt(13)       
     
    odoc.tables[3].cell(1,5).paragraphs[0].add_run(df['Htd'].iloc[0]).font.size = Pt(13)
    odoc.tables[3].cell(2,5).paragraphs[0].add_run(df['Htd'].iloc[1]).font.size = Pt(13) 
    odoc.tables[3].cell(3,5).paragraphs[0].add_run(df['Htd'].iloc[2]).font.size = Pt(13) 
    odoc.tables[3].cell(4,5).paragraphs[0].add_run(df['Htd'].iloc[3]).font.size = Pt(13)    
    
    odoc.tables[4].cell(1,1).paragraphs[0].add_run(df['Htdtk'].iloc[0]).font.size = Pt(13)
    odoc.tables[4].cell(1,2).paragraphs[0].add_run(df['Htdtk'].iloc[1]).font.size = Pt(13) 
    odoc.tables[4].cell(1,3).paragraphs[0].add_run(df['Htdtk'].iloc[2]).font.size = Pt(13) 
    odoc.tables[4].cell(1,4).paragraphs[0].add_run(df['Htdtk'].iloc[3]).font.size = Pt(13)    
    
    
    odoc.tables[6].cell(1,1).paragraphs[0].add_run(lamtron_Q(df['qdtdk'].iloc[4])).font.size = Pt(13)
    odoc.tables[6].cell(2,1).paragraphs[0].add_run(lamtron_Q(df['qdtdk'].iloc[5])).font.size = Pt(13) 
    odoc.tables[6].cell(3,1).paragraphs[0].add_run(lamtron_Q(df['qdtdk'].iloc[6])).font.size = Pt(13) 
    odoc.tables[6].cell(4,1).paragraphs[0].add_run(lamtron_Q(df['qdtdk'].iloc[7])).font.size = Pt(13) 
    
    odoc.tables[6].cell(1,2).paragraphs[0].add_run(lamtron_Q(df['qxdk'].iloc[4])).font.size = Pt(13)
    odoc.tables[6].cell(2,2).paragraphs[0].add_run(lamtron_Q(df['qxdk'].iloc[5])).font.size = Pt(13) 
    odoc.tables[6].cell(3,2).paragraphs[0].add_run(lamtron_Q(df['qxdk'].iloc[6])).font.size = Pt(13) 
    odoc.tables[6].cell(4,2).paragraphs[0].add_run(lamtron_Q(df['qxdk'].iloc[7])).font.size = Pt(13) 
    
    odoc.tables[6].cell(1,3).paragraphs[0].add_run(lamtron_Q(df['qcmdk'].iloc[4])).font.size = Pt(13)
    odoc.tables[6].cell(2,3).paragraphs[0].add_run(lamtron_Q(df['qcmdk'].iloc[5])).font.size = Pt(13) 
    odoc.tables[6].cell(3,3).paragraphs[0].add_run(lamtron_Q(df['qcmdk'].iloc[6])).font.size = Pt(13) 
    odoc.tables[6].cell(4,3).paragraphs[0].add_run(lamtron_Q(df['qcmdk'].iloc[7])).font.size = Pt(13)     
    
    odoc.tables[6].cell(1,4).paragraphs[0].add_run(lamtron_Q(df['qdb'].iloc[4])).font.size = Pt(13)
    odoc.tables[6].cell(2,4).paragraphs[0].add_run(lamtron_Q(df['qdb'].iloc[5])).font.size = Pt(13) 
    odoc.tables[6].cell(3,4).paragraphs[0].add_run(lamtron_Q(df['qdb'].iloc[6])).font.size = Pt(13) 
    odoc.tables[6].cell(4,4).paragraphs[0].add_run(lamtron_Q(df['qdb'].iloc[7])).font.size = Pt(13)       
     
    odoc.tables[6].cell(1,5).paragraphs[0].add_run(lamtron_Q(df['Qmax'].iloc[4])).font.size = Pt(13)
    odoc.tables[6].cell(2,5).paragraphs[0].add_run(lamtron_Q(df['Qmax'].iloc[5])).font.size = Pt(13) 
    odoc.tables[6].cell(3,5).paragraphs[0].add_run(lamtron_Q(df['Qmax'].iloc[6])).font.size = Pt(13) 
    odoc.tables[6].cell(4,5).paragraphs[0].add_run(lamtron_Q(df['Qmax'].iloc[7])).font.size = Pt(13)  
     
    odoc.tables[6].cell(1,6).paragraphs[0].add_run(df['Hdb'].iloc[4]).font.size = Pt(13)
    odoc.tables[6].cell(2,6).paragraphs[0].add_run(df['Hdb'].iloc[5]).font.size = Pt(13) 
    odoc.tables[6].cell(3,6).paragraphs[0].add_run(df['Hdb'].iloc[6]).font.size = Pt(13) 
    odoc.tables[6].cell(4,6).paragraphs[0].add_run(df['Hdb'].iloc[7]).font.size = Pt(13)    
    
    odoc.tables[6].cell(1,7).paragraphs[0].add_run(df['Hmax'].iloc[4]).font.size = Pt(13)
    odoc.tables[6].cell(2,7).paragraphs[0].add_run(df['Hmax'].iloc[5]).font.size = Pt(13) 
    odoc.tables[6].cell(3,7).paragraphs[0].add_run(df['Hmax'].iloc[6]).font.size = Pt(13) 
    odoc.tables[6].cell(4,7).paragraphs[0].add_run(df['Hmax'].iloc[7]).font.size = Pt(13)   
       
    for t in range(1,7):
        for row in odoc.tables[t].rows[1:]:
            for cell in row.cells[1:]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    
    odoc.save(pth_docx)
    # convert(pth_docx,pth_docx.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
def tin_lulu_load1():
    now = datetime.now()
    tgpt = thoigianphattin()
    pth_docx = read_txt('path_tin/LULU.txt') + '/DHC_LULU_' + tgpt.strftime('%Y%m%d_%H%M') + '.docx'
    odoc = Document(pth_docx)
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    excel = client.Dispatch("Excel.Application")
    excel.Visible = True
    book = excel.Workbooks.Open(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx')
    
    ttttt =0
    for i in range(3,1000):
        dlngay = book.Worksheets('DRHN').Cells(i,1).Value
        dlngay = dlngay.Format('%Y-%m-%d %H:%M:%S')
        dlngay = datetime.strptime(dlngay,'%Y-%m-%d %H:%M:%S')
        
        if tgpt.hour ==4 or tgpt.hour ==10 or tgpt.hour ==16 or tgpt.hour ==22:
            ttttt =1
            tgpt = tgpt +timedelta(hours=2.5)
            
        if dlngay.strftime('%Y%m%d %H') == tgpt.strftime('%Y%m%d %H'):
            if ttttt==1:
                book.Worksheets('DRHN').Cells(i,10).Value = "=(c" + str(i) + "-f" + str(i) + "-h" + str(i) + ")*108/10^4"
            else:
                book.Worksheets('DRHN').Cells(i,10).Value = "=(c" + str(i) + "-f" + str(i) + "-h" + str(i) + ")*216/10^4"
            break
        
    for t in range(3,7,3):
        for row in odoc.tables[t].rows[1:]:
            for cell in row.cells[1:]:
                cell.text = ''
    for row in odoc.tables[4].rows[1:]:
        for cell in row.cells[1:]:
            cell.text = ''
                # for paragraph in cell.paragraphs:
                #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    
    # qtb về hồ
    odoc.tables[3].cell(1,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-3,2).Value)).font.size = Pt(13)
    odoc.tables[3].cell(2,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-2,2).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(3,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-1,2).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(4,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i,2).Value)).font.size = Pt(13) 
    
    #Q xả duy trì
    odoc.tables[3].cell(1,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-3,7).Value)).font.size = Pt(13)
    odoc.tables[3].cell(2,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-2,7).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(3,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-1,7).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(4,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i,7).Value)).font.size = Pt(13) 
    
    # Q xa tran
    odoc.tables[3].cell(1,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-3,5).Value)).font.size = Pt(13)
    odoc.tables[3].cell(2,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-2,5).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(3,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-1,5).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(4,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i,5).Value)).font.size = Pt(13)     
    
    # Q chay may
    odoc.tables[3].cell(1,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-3,9).Value)).font.size = Pt(13)
    odoc.tables[3].cell(2,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-2,9).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(3,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i-1,9).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(4,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i,9).Value)).font.size = Pt(13)       
    
    # H hồ
    odoc.tables[3].cell(1,5).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-3,12).Value)).font.size = Pt(13)
    odoc.tables[3].cell(2,5).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-2,12).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(3,5).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-1,12).Value)).font.size = Pt(13) 
    odoc.tables[3].cell(4,5).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i,12).Value)).font.size = Pt(13)    
    
    # muc nuoc tra khuc
    odoc.tables[4].cell(1,1).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-3,21).Value)).font.size = Pt(13)
    odoc.tables[4].cell(1,2).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-2,21).Value)).font.size = Pt(13) 
    odoc.tables[4].cell(1,3).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i-1,21).Value)).font.size = Pt(13) 
    odoc.tables[4].cell(1,4).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i,21).Value)).font.size = Pt(13)    
    
    #Q xả duy trì dk
    odoc.tables[6].cell(1,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+1,6).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+2,6).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+3,6).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,1).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+4,6).Value)).font.size = Pt(13) 
    
    #Q xả tràn dk
    odoc.tables[6].cell(1,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+1,4).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+2,4).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+3,4).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,2).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+4,4).Value)).font.size = Pt(13) 
    
    #Q chay may dk
    odoc.tables[6].cell(1,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+1,8).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+2,8).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+3,8).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,3).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+4,8).Value)).font.size = Pt(13)     
    
    # Qtb ve ho db
    odoc.tables[6].cell(1,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+1,3).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+2,3).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+3,3).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,4).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+4,3).Value)).font.size = Pt(13)     
      
    # Q max ve ho
    odoc.tables[6].cell(1,5).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+1,19).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,5).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+2,19).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,5).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+3,19).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,5).paragraphs[0].add_run(lamtron_Q(book.Worksheets('DRHN').Cells(i+4,19).Value)).font.size = Pt(13)  
     
     #H hồ db
    odoc.tables[6].cell(1,6).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+1,13).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,6).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+2,13).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,6).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+3,13).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,6).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+4,13).Value)).font.size = Pt(13)    
    # H max tra khuc
    odoc.tables[6].cell(1,7).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+1,20).Value)).font.size = Pt(13)
    odoc.tables[6].cell(2,7).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+2,20).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(3,7).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+3,20).Value)).font.size = Pt(13) 
    odoc.tables[6].cell(4,7).paragraphs[0].add_run('{:.2f}'.format(book.Worksheets('DRHN').Cells(i+4,20).Value)).font.size = Pt(13)   
       
    for t in range(1,7):
        for row in odoc.tables[t].rows[1:]:
            for cell in row.cells[1:]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    
    odoc.save(pth_docx)
    
    # book.SaveAs(r"C:\Users\Administrator\Desktop\chep so.xlsx")
    book.Save()
    book.Close()
    excel.Quit()
    # convert(pth_docx,pth_docx.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')