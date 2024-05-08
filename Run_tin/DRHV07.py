from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
from func.load_data import mo_word
from func.CDH_TTB_API import TTB_API_mucnuoc10day
import pyodbc

def custom_round(value):
    if value != '-' and value % 1 == 0.5:
        return round(value + 0.1,0)
    return int(round(value,0))


selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value


def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def sobt():
    pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    # print(pth)
    now = datetime.now()
    if now.strftime('%Y%m%d') in pth:
        os.remove(pth)
        messagebox.showinfo('Thong bao','Đã xóa file tồn tại' + pth.split('\\')[-1])
        pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    
    odoc = Document(pth)
    for a in odoc.tables[0].cell(0,0).paragraphs:
        if 'Số' in a.text:
            dl = str(a.text)
            sbt = dl[dl.index('-')+1:dl.index('/')]
    return int(sbt) + 1

def xacdinhngaydb():
    now = datetime.now()
    for a in range(3,10):
        tttt = now + timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def xacdinhngaydaqua():
    now = datetime.now()
    for a in range(3,10):
        tttt = now - timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def tin_nenKT_05day():
    now = datetime.now()
    odoc = Document('TINMAU/ST_TVHV07.docx')
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    # so ban tin
    ngaydb = datetime(now.year,9,1)
    delta = now - ngaydb
    delta = datetime(2024,9,1) - ngaydb
    sbt = int((delta.days+1)/7) + 1
    # print(sbt)
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số: '+ str(sbt) + '/DBHVST2-ĐQNAM'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Nam' in dl:
                pr.text=''
                ntn = 'Quảng Nam, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
            
    
    # ngaydb = now + timedelta(days=6)
    # ngaydb = ngaydb - timedelta(days=1)
    # ngaydb = ngaydb.strftime('%d/%m/%Y')
    # ngaytd = xacdinhngaydaqua()
    # ngaytd = ngaytd.strftime('%d/%m/%Y')
    
    
    for pr in odoc.paragraphs:
        dl = pr.text
        if 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN VỪA' in dl:
            # ban tin tiep theo
            ntn = 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN VỪA'
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(14)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '(Từ ngày 12 đến 18/12/2023)' in dl:
            # ban tin tiep theo
            ntn = '(Từ ngày {} đến ngày {})'.format(now.strftime('%d/%m'),(now + timedelta(days=6)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Bảng 1' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 1: Tổng hợp lượng mưa (mm) từ ngày {} - {}'.format((now - timedelta(days=6)).strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Bảng 2' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 2: Đặc trưng lưu lượng (m3/s) từ ngày {} - {}'.format((now - timedelta(days=6)).strftime('%d/%m/%Y'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        elif 'Bảng 3' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 3: Dự báo lượng mưa (mm) từ ngày {} đến ngày {}'.format(now .strftime('%d/%m'),(now + timedelta(days=6)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13) 
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        elif  'Từ ngày 12 - 16/12, dòng chảy biến đổi chậm' in dl:
            # ban tin tiep theo
            ntn = 'Từ ngày {} - {}, dòng chảy về hồ Thuỷ điện Sông Tranh 2 tiếp tục biến đổi chậm.'.format(now.strftime('%d/%m/%Y'),(now + timedelta(days=6)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
        elif 'Bảng 4' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 4: Đặc trưng lưu lượng (m3/s) dự báo từ ngày {} đến ngày {}'.format(now.strftime('%d/%m'),(now + timedelta(days=6)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        elif 'Dự báo viên' in dl:
            # ban tin tiep theo
            ntn = 'Dự báo viên: {}'.format(selected_value)
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)  
        elif 'Thời gian ban hành bản tin tiếp theo' in dl:
            # ban tin tiep theo
            ntn = 'Thời gian ban hành bản tin tiếp theo 8 giờ 00 phút ngày {}'.format((now + timedelta(days=6)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)              
        elif 'Phụ lục: Tổng hợp lượng mưa các trạm đo trên lưu vực Sông Tranh' in dl:
            # ban tin tiep theo
            ntn = 'Phụ lục: Tổng hợp lượng mưa các trạm đo trên lưu vực Sông Tranh từ {} - {}'.format((now-timedelta(days=6)).strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        

    # lay so lieu mua
    pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA.accdb'
    # pth25 = r'D:\PM_PYTHON\SONGTRANH\DATA\DATA.accdb'
    FileName=(pth25)
    cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
    query = "SELECT * FROM mua"
    mua = pd.read_sql(query, cnxn)
    bd_mua = datetime((now-timedelta(days=7)).year,(now-timedelta(days=7)).month,(now-timedelta(days=7)).day,20)
    # print(bd_mua)
    # print(datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,19))
    mua = mua[(mua['thoigian'] >=bd_mua) & (mua['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,19))]
    mua.set_index('thoigian',inplace=True)
    mua = mua.replace('-',0)
    # print(mua)
    mua = mua.astype(float)

    mua10 = mua.sum()
    mua10 = mua10.replace(0.0,'-')
    mua10 = mua10.map(lambda x: custom_round(float(x)) if x != '-' else '-')
    dt_mua = mua10.replace('-',0)
    max1 = max( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    min1 = min( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    max2 = max( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    min2 = min( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    max3 = max( dt_mua['trabui'], dt_mua['dapsongtranh'])
    min3 = min( dt_mua['trabui'], dt_mua['dapsongtranh'])
    
    # mua ngay de nhan xet
    muangay = mua.rolling(24,min_periods=1).sum()
    muangay = muangay[muangay.index.hour==19]
    muangay = muangay.applymap("{0:.1f}".format)

    # Thêm dữ liệu từ DataFrame vào bảng
    for a,i in enumerate(range(muangay.shape[0])):
        odoc.tables[1].cell(i+1, 0).text = (bd_mua + timedelta(days=a+1)).strftime('%d/%m')
        for j in range(muangay.shape[-1]):
            odoc.tables[1].cell(i+1, j+1).text = str(muangay.values[i, j])

    # dieu chinh lai format cho bang thong ke
    for a in range(0,12):
        for b in range(0,12):
            try:
                for pr in odoc.tables[1].cell(a,b).paragraphs:
                    pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in pr.runs:
                        run.font.size = Pt(8)
                        run.bold = False  
            except:
                pass
    
    # print(muangay)

        
    mua10 = mua10.astype(str)
    # print(mua10)
    # bang top hop mua so 1   
    odoc.tables[2].cell(1,0).text= 'Từ {} - {}'.format((bd_mua+ timedelta(days=1)).strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
    odoc.tables[2].cell(1,1).text= str(min1) + ' - ' + str(max1)
    odoc.tables[2].cell(1,2).text= str(min2) + ' - ' + str(max2)
    odoc.tables[2].cell(1,3).text= str(min3) + ' - ' + str(max3)   
    
    for j in range(0,4):
        pr = odoc.tables[2].cell(1,j).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # bang du bao mua 3
    odoc.tables[4].cell(1,0).text= 'Từ {} - {}'.format(now.strftime('%d/%m'),ngaydb.strftime('%d/%m/%Y'))
    odoc.tables[4].cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # bang top hop mua so 6
    odoc.tables[7].cell(1,2).text= mua10['tralinh']
    odoc.tables[7].cell(2,2).text= mua10['tranam2']
    odoc.tables[7].cell(3,2).text= mua10['travan']
    odoc.tables[7].cell(4,2).text= mua10['tracang']
    odoc.tables[7].cell(5,2).text= mua10['tramai']
    odoc.tables[7].cell(6,2).text= mua10['tragiac']
    odoc.tables[7].cell(7,2).text= mua10['tradon']
    odoc.tables[7].cell(8,2).text= mua10['traleng']
    odoc.tables[7].cell(9,2).text= mua10['dapsongtranh']
    odoc.tables[7].cell(10,2).text= mua10['trabui']

    for j in range(1,11):
        pr = odoc.tables[7].cell(j,2).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    for j in range(1,11):
        pr = odoc.tables[7].cell(j,2).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # print(mua10)

    try:
        mucnuoc = TTB_API_mucnuoc10day()
        mucnuoc = mucnuoc.reset_index(False)
        mucnuoc.rename(columns={'time':'thoigian','qden':'qdenho'},inplace=True)
        cnxn.close()
    except:
        query = "SELECT thoigian,qdenho FROM thuyvan"
        mucnuoc = pd.read_sql(query, cnxn)
        cnxn.close()
    # print(bd_mua + timedelta(hours=4))
    mucnuoc = mucnuoc[(mucnuoc['thoigian'] >=(bd_mua + timedelta(hours=4))) & (mucnuoc['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,23))]

    mucnuoc.set_index('thoigian',inplace=True)
    # print(mucnuoc)
    # mucnuoc.to_excel('kiemtrah.xlsx')
    mucnuoc = mucnuoc.astype(float)
    # bang thuc do luu luong 2
    odoc.tables[3].cell(1,1).text= '{0:.1f}'.format(mucnuoc['qdenho'].mean())
    odoc.tables[3].cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # # print(mua6h)
    pth = read_txt('path_tin/DRHV.txt') + '/QNAM_BT07_STRANH_{}.docx'.format(now.strftime('%Y%m%d')) 
    odoc.save(pth)
    # # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    