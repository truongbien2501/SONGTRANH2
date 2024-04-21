from tkinter.font import ITALIC
import pandas as pd
from openpyxl import load_workbook,worksheet
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Pt,Inches
from docx2pdf import convert
from tkinter import messagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from func.Seach_file import tim_file,read_txt
from func.load_data import mo_word
import pyodbc
selected_value = None
def set_selected_value(value):
    global selected_value
    selected_value = value

duyettin = None

def custom_round(value):
    if value != '-' and value % 1 == 0.5:
        return round(value + 0.1,0)
    return int(round(value,0))

def set_selected_duyet(value):
    global duyettin
    duyettin = value

def lamtron_Q(q):
    q = float(q)
    if q < 10:
        luuluong = '{:.2f}'.format(q)
    elif q >=10 and q <100:
        luuluong = '{:.1f}'.format(q)
    elif q >=100:
        luuluong = '{:.0f}'.format(q)
    else:
        luuluong = ''
    return str(luuluong )

def sobt():
    pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    # print(pth)
    now = datetime.now()
    if now.strftime('%Y%m%d') in pth:
        os.remove(pth)
        messagebox.showinfo('Thong bao','Đã xóa file tồn tại' + pth.split('\\')[-1])
        pth = tim_file(read_txt('path_tin/DRHV.txt'),'.docx')
    
    odoc = Document(pth)
    for a in odoc.tables[0].cell(0,0).paragraphs:
        if 'Số' in a.text:
            dl = str(a.text)
            sbt = dl[dl.index(':')+1:dl.index('/')]
    return int(sbt) + 1

def xacdinhngaydb():
    now = datetime.now()
    for a in range(5,12):
        tttt = now + timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def xacdinhngaydaqua():
    now = datetime.now()
    for a in range(5,12):
        tttt = now - timedelta(days=a)
        if tttt.strftime('%d')[-1]=='1' and ('3' not in tttt.strftime('%d')) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break
    return ngay

def tin_nenKT_10day():
    now = datetime.now()
    odoc = Document('TINMAU/ST_TVHV10.docx')
    # odoc = Document(r'D:\PM_PYTHON\SONGTRANH\TINMAU\ST_TVHV10.docx')
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    # so ban tin
    sbt = sobt()
    # sbt = 25
    for t in range(0,2):
        for pr in odoc.tables[0].cell(0,t).paragraphs:
            dl = pr.text
            if 'Số:' in dl:
                pr.text=''
                soso = 'Số: '+ str(sbt) + '/DBHVST2-ĐQNAM'
                run = pr.add_run(soso)
                run.bold = False
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'Quảng Nam' in dl:
                pr.text=''
                ntn = 'Quảng Nam, ngày ' + now.strftime('%d') + ' tháng ' + now.strftime('%m') + ' năm ' + now.strftime('%Y')
                run = pr.add_run(ntn)
                run.italic = True
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break
        for run in pr.runs:
            font = run.font
            font.name = 'Times New Roman'
            
    
    ngaydb = xacdinhngaydb()
    ngaydb = ngaydb - timedelta(days=1)    
    ngaytd = xacdinhngaydaqua()
    bd_mua = datetime(ngaytd.year,ngaytd.month,ngaytd.day,20)
    
    # ngaydb = datetime.now()
    # ngaytd = ngaydb - timedelta(days=10)
    # bd_mua = datetime(ngaytd.year,ngaytd.month,ngaytd.day,0)
    
    for pr in odoc.paragraphs:
        dl = pr.text
        if 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN VỪA' in dl:
            # ban tin tiep theo
            ntn = 'BẢN TIN DỰ BÁO THỜI TIẾT THUỶ VĂN HẠN VỪA'
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(14)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif '(Từ ngày 21 đến 31/3/2024)' in dl:
            # ban tin tiep theo
            ntn = '(Từ ngày {} đến ngày {})'.format(now.strftime('%d/%m/'),ngaydb.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Bảng 1: Tổng hợp lượng mưa (mm) từ ngày 11 - 20/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 1: Tổng hợp lượng mưa (mm) từ ngày {} - {}'.format(ngaytd.strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif 'Bảng 2: Đặc trưng lưu lượng (m3/s) từ ngày 11 đến ngày 20/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 2: Đặc trưng lưu lượng (m3/s) từ ngày {} - {}'.format(ngaytd.strftime('%d/%m/%Y'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        elif 'Bảng 3: Dự báo lượng mưa (mm) từ ngày 21 - 31/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 3: Dự báo lượng mưa (mm) từ ngày {} đến ngày {}'.format(now.strftime('%d/%m'),ngaydb.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13) 
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        elif  'Từ ngày 21 - 31/3/2024, dòng chảy về hồ Thuỷ điện Sông Tranh 2 tiếp tục biến đổi chậm.' in dl:
            # ban tin tiep theo
            ntn = 'Từ ngày {} - {}, dòng chảy về hồ Thuỷ điện Sông Tranh 2 tiếp tục biến đổi chậm.'.format(now.strftime('%d/%m/%Y'),ngaydb.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.font.size = Pt(13)
        elif 'Bảng 4: Đặc trưng lưu lượng (m3/s) dự báo từ ngày 21 đến ngày 31/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Bảng 4: Đặc trưng lưu lượng (m3/s) dự báo từ ngày {} đến ngày {}'.format(now.strftime('%d/%m'),ngaydb.strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        elif 'Dự báo viên' in dl:
            # ban tin tiep theo
            ntn = 'Dự báo viên: {}'.format(selected_value)
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.italic = True
            run.font.size = Pt(13)  
        elif 'Tổng hợp lượng mưa các trạm đo trên lưu vực Sông Tranh từ 11 - 20/3/2024' in dl:
            # ban tin tiep theo
            ntn = 'Tổng hợp lượng mưa các trạm đo trên lưu vực Sông Tranh từ {} - {}'.format(ngaytd.strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
            pr.text  =''
            run = pr.add_run(ntn)
            # run.bold = True
            run.font.size = Pt(13)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        

    # lay so lieu mua
    pth25 = read_txt('path_tin/DATA_EXCEL.txt') + '/QNAM.accdb'
    # pth25 = r'D:\PM_PYTHON\SONGTRANH\DATA\QNAM.accdb'
    FileName=(pth25)
    cnxn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + FileName + ';')
    query = "SELECT * FROM mua"
    mua = pd.read_sql(query, cnxn)
    
    mua = mua[(mua['thoigian'] >=datetime((bd_mua-timedelta(days=1)).year,(bd_mua-timedelta(days=1)).month,(bd_mua-timedelta(days=1)).day,20)) & (mua['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,19))]
    mua.set_index('thoigian',inplace=True)
    mua = mua.astype(float)
    mua10 = mua.sum()
    mua10 = mua10.replace(0.0,'-')
    mua10 = mua10.map(lambda x: custom_round(float(x)) if x != '-' else '-')
    dt_mua = mua10.replace('-',0)
    max1 = max( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    min1 = min( dt_mua['tralinh'], dt_mua['tranam2'], dt_mua['travan'], dt_mua['tracang'],dt_mua['tramai'])
    max2 = max( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    min2 = min( dt_mua['tragiac'], dt_mua['tradon'], dt_mua['traleng'])
    max3 = max( dt_mua['trabui'], dt_mua['dapsongtranh'])
    min3 = min( dt_mua['trabui'], dt_mua['dapsongtranh'])
    
    # mua ngay de nhan xet
    muangay = mua.rolling(24,min_periods=1).sum()
    muangay = muangay[muangay.index.hour==19]
    muangay = muangay.applymap("{0:.1f}".format)

    # Thêm dữ liệu từ DataFrame vào bảng
    for a,i in enumerate(range(muangay.shape[0])):
        odoc.tables[1].cell(i+1, 0).text = (bd_mua + timedelta(days=a)).strftime('%d/%m/')
        for j in range(muangay.shape[-1]):
            odoc.tables[1].cell(i+1, j+1).text = str(muangay.values[i, j])

    # dieu chinh lai format cho bang thong ke
    for a in range(0,12):
        for b in range(0,12):
            try:
                for pr in odoc.tables[1].cell(a,b).paragraphs:
                    pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in pr.runs:
                        run.font.size = Pt(8)
                        run.bold = False  
            except:
                pass
    
    # print(muangay)

        
    mua10 = mua10.astype(str)
    # print(mua10)
    # bang top hop mua so 1   
    odoc.tables[2].cell(1,0).text= 'Từ {} - {}'.format(bd_mua.strftime('%d/%m'),(now-timedelta(days=1)).strftime('%d/%m/%Y'))
    odoc.tables[2].cell(1,1).text= str(min1) + ' - ' + str(max1)
    odoc.tables[2].cell(1,2).text= str(min2) + ' - ' + str(max2)
    odoc.tables[2].cell(1,3).text= str(min3) + ' - ' + str(max3)   
    
    for j in range(0,4):
        pr = odoc.tables[1].cell(1,j).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # bang du bao mua 3
    odoc.tables[4].cell(1,0).text= 'Từ {} - {}'.format(now.strftime('%d/%m'),ngaydb.strftime('%d/%m/%Y'))
    odoc.tables[4].cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # bang top hop mua so 6
    odoc.tables[7].cell(1,1).text= mua10['tralinh']
    odoc.tables[7].cell(2,1).text= mua10['tranam2']
    odoc.tables[7].cell(3,1).text= mua10['travan']
    odoc.tables[7].cell(4,1).text= mua10['tracang']
    odoc.tables[7].cell(5,1).text= mua10['tramai']
    odoc.tables[7].cell(6,1).text= mua10['tragiac']
    odoc.tables[7].cell(7,1).text= mua10['tradon']
    odoc.tables[7].cell(8,1).text= mua10['traleng']
    odoc.tables[7].cell(9,1).text= mua10['dapsongtranh']
    odoc.tables[7].cell(10,1).text= mua10['trabui']

    for j in range(1,11):
        pr = odoc.tables[7].cell(j,1).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    for j in range(1,11):
        pr = odoc.tables[7].cell(j,1).paragraphs[0]
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # print(mua10)

    query = "SELECT thoigian,qdenho FROM thuyvan"
    mucnuoc = pd.read_sql(query, cnxn)
    cnxn.close()
    mucnuoc = mucnuoc[(mucnuoc['thoigian'] >=bd_mua) & (mucnuoc['thoigian'] <= datetime((now-timedelta(days=1)).year,(now-timedelta(days=1)).month,(now-timedelta(days=1)).day,23))]
    mucnuoc.set_index('thoigian',inplace=True)
    mucnuoc = mucnuoc.astype(float)
    # bang thuc do luu luong 2
    odoc.tables[3].cell(1,1).text= '{0:.1f}'.format(mucnuoc['qdenho'].mean())
    odoc.tables[3].cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # # print(mua6h)
    pth = read_txt('path_tin/DRHV.txt') + '/QNAM_BT10_STRANH_{}.docx'.format(now.strftime('%Y%m%d')) 
    odoc.save(pth)
    # # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
# tin_nenKT_10day()   
    
    
def tin_tv10_load():
    now = datetime.now()
    pth = read_txt('path_tin/DRHV.txt') + '/DHC_TV05_' + now.strftime('%Y%m%d_1600') + '.docx'
    odoc = Document(pth)
    style = odoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
   
    df = pd.read_excel(read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx',engine="openpyxl",sheet_name='DRHV05')
    
    df["Qtb_td"] = df["Qtb_td"].apply(lambda x: f"{x}")
    df["Qtb_td"] = df["Qtb_td"].astype(float)
    df["Qtb_td"] = df["Qtb_td"].map('{0:.1f}'.format)
    df["W"] = df["W"].apply(lambda x: f"{x}")
    df["W"] = df["W"].astype(float)
    df["W"] = df["W"].map('{0:.1f}'.format)
    
    df["Qtháng"] = df["Qtháng"].apply(lambda x: f"{x}")
    df["Qtháng"] = df["Qtháng"].astype(str)
    # df["Qtháng"] = df["Qtháng"].map('{0:.1f}'.format)
    df["W tháng"] = df["W tháng"].apply(lambda x: f"{x}")
    # df["Ngày"] = df["Ngày"].apply(lambda x: f"{x}")
    df["W tháng"] = df["W tháng"].astype(str)
    # df["W tháng"] = df["W tháng"].map('{0:.1f}'.format)
    # print(df.columns)
    
    
    now = datetime(now.year,now.month,now.day)
    kt = xacdinhngaydb()
    bd = xacdinhngaydaqua()

    df['time'] = pd.date_range(start=datetime(now.year,9,1), periods=len(df['Ngày']), freq="D")
    qtb = df[df['time']==now]['Qtb_td']
    tongluong = df[df['time']==now]['W']
    odoc.tables[2].cell(0,1).text = qtb
    odoc.tables[2].cell(1,1).text = tongluong
    
    
    df_db = df.loc[(df['time'] >= bd) & (df['time'] < now) ] # so lieu thuc 

    
    for i in range(1,6):  
        for j in range(1,3):         
            odoc.tables[4].cell(i,j).text = ''
    
    
    
    odoc.tables[4].cell(1,1).text = df[df['time']==now]['Qtháng']
    odoc.tables[4].cell(2,1).text = df[df['time']==(now + timedelta(days=1))]['Qtháng']
    odoc.tables[4].cell(3,1).text = df[df['time']==(now + timedelta(days=2))]['Qtháng']
    odoc.tables[4].cell(4,1).text = df[df['time']==(now + timedelta(days=3))]['Qtháng']
    odoc.tables[4].cell(5,1).text = df[df['time']==(now + timedelta(days=4))]['Qtháng']
    
    odoc.tables[4].cell(1,2).text = df[df['time']==now]['W tháng']
    odoc.tables[4].cell(2,2).text = df[df['time']==(now + timedelta(days=1))]['W tháng']
    odoc.tables[4].cell(3,2).text = df[df['time']==(now + timedelta(days=2))]['W tháng']
    odoc.tables[4].cell(4,2).text = df[df['time']==(now + timedelta(days=3))]['W tháng']
    odoc.tables[4].cell(5,2).text = df[df['time']==(now + timedelta(days=4))]['W tháng']
    
    ngaydb = xacdinhngaydb()
    num_days = abs(ngaydb - now).days
    if num_days ==6:
        odoc.tables[4].cell(6,1).text = df[df['time']==(now + timedelta(days=5))]['Qtháng']
        odoc.tables[4].cell(6,2).text = df[df['time']==(now + timedelta(days=5))]['W tháng']
    
    for i in range(1,6):  
        for j in range(0,3):         
            pr = odoc.tables[4].cell(i,j).paragraphs[0]
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    odoc.save(pth)
    mo_word(pth)
    # convert(pth,pth.replace('.docx','.pdf'))
    messagebox.showinfo('Thông báo','OK!')
    
    