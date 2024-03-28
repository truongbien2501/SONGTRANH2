from email.header import decode_header
import os
from imbox import Imbox
from tkinter import messagebox
from selenium import webdriver
import time
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import Select
import pandas as pd
from datetime import datetime,timedelta
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
import numpy as np
from docx import Document
from win32com import client
from func.Seach_file import tim_file,read_txt
import os
def mo_excel():
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    # pth = os.getcwd() + '/DATA/DR_THUYVAN.xlsx'
    excel = client.Dispatch("Excel.Application")
    excel.Visible = True
    obook = excel.Workbooks.Open(pth)
    obook.Worksheets(1).Select()
def mo_word(pth):
    word = client.Dispatch("Word.Application")
    word.Visible = True
    odoc = word.Documents.Open(pth)

def downloadattmail():
    host = "imap.gmail.com"
    username = "kttvqngai@gmail.com"
    password = 'dhwwemyolidvvxuu'
    download_folder = os.getcwd()

    if not os.path.isdir(download_folder):
        os.makedirs(download_folder, exist_ok=True)
        
    mail = Imbox(host, username=username, password=password, ssl=True, ssl_context=None, starttls=False)
    messages = mail.messages(sent_from='pxvhdakdrinh@gmail.com') # defaults to inbox
    for (uid, message) in messages[-1:]:
        mail.mark_seen(uid) # optional, mark message as read

        for idx, attachment in enumerate(message.attachments):
            try:
                att_fn = attachment.get('filename')
                if 'Đak' in att_fn:
                    # download_path = f"{download_folder}/{att_fn}"
                    with open('SOLIEU/DAKDRINH.doc', "wb") as fp:
                        fp.write(attachment.get('content').read())
            except:
                print('traceback.print_exc()')
    mail.logout()
    convertdocx()
    messagebox.showinfo('Thông báo', 'OK')
    
def downloadattmail_lulu():
    host = "imap.gmail.com"
    username = "kttvqngai@gmail.com"
    password = 'dhwwemyolidvvxuu'
    download_folder = os.getcwd()

    if not os.path.isdir(download_folder):
        os.makedirs(download_folder, exist_ok=True)
        
    mail = Imbox(host, username=username, password=password, ssl=True, ssl_context=None, starttls=False)
    messages = mail.messages(sent_from='pxvhdakdrinh@gmail.com') # defaults to inbox
    for (uid, message) in messages[-2:]:
        mail.mark_seen(uid) # optional, mark message as read

        for idx, attachment in enumerate(message.attachments):
            try:
                att_fn = attachment.get('filename')
                # print(att_fn)
                if '3h KTTV' in att_fn:
                    # download_path = f"{download_folder}/{att_fn}"
                    with open('SOLIEU/DAKDRINH_LULU.doc', "wb") as fp:
                        fp.write(attachment.get('content').read())
            except:
                print('traceback.print_exc()')
    mail.logout()
    convertdocx_lulu()
    messagebox.showinfo('Thông báo', 'OK')
# downloadattmail()

def loadmua():
    driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()))
    now = datetime.now()
    bd = now - timedelta(hours=23)
    pth = 'http://hothuydien.kttvttb.vn/dakrinh/#'
    driver.maximize_window()
    driver.get(pth)
    time.sleep(5)
    element = driver.find_element(by=By.CSS_SELECTOR,value='#mn2 > a')
    element.click()
    time.sleep(2)
    
    #dakdrinh
    select_element = Select(driver.find_element(by=By.CSS_SELECTOR,value='#chontram'))
    select_element.select_by_value('1')
    time.sleep(2)
    element = driver.find_element(by=By.CSS_SELECTOR,value='#menu2 > div:nth-child(3) > button.btn.btn-default')
    element.click()
    time.sleep(5)
    table_element = driver.find_element(by=By.CSS_SELECTOR, value= "#data2 > table")
    table_header = []
    for th in table_element.find_elements(by=By.TAG_NAME, value="th"):
        table_header.append(th.text)

    table_data = []
    for row in table_element.find_elements(by=By.TAG_NAME,value="tr"):
        row_data = []
        for cell in row.find_elements(by=By.TAG_NAME,value="td"):
            row_data.append(cell.text)
        table_data.append(row_data)
    daumoi = pd.DataFrame(table_data,columns=table_header)
    daumoi = daumoi.T
    daumoi = daumoi.iloc[4:,:]
    tg_bd = daumoi.index[0]
    thoigian = datetime(bd.year,int(tg_bd[-2:]),int(tg_bd[6:8]),int(tg_bd[:2]),int(tg_bd[3:5]))
    daumoi.insert(0, 'time', pd.date_range(thoigian,periods=len(daumoi.index),freq='H'))
    daumoi = daumoi.sort_values(by='time')
    daumoi = daumoi.drop(columns=[0])
    daumoi.rename(columns={1:'daumoi'},inplace=True)
    # daltang
    select_element.select_by_value('3')
    time.sleep(2)
    element = driver.find_element(by=By.CSS_SELECTOR,value='#menu2 > div:nth-child(3) > button.btn.btn-default')
    element.click()
    time.sleep(5)
    table_element = driver.find_element(by=By.CSS_SELECTOR, value= "#data2 > table")
    table_header = []
    for th in table_element.find_elements(by=By.TAG_NAME, value="th"):
        table_header.append(th.text)

    table_data = []
    for row in table_element.find_elements(by=By.TAG_NAME,value="tr"):
        row_data = []
        for cell in row.find_elements(by=By.TAG_NAME,value="td"):
            row_data.append(cell.text)
        table_data.append(row_data)
    daktang = pd.DataFrame(table_data,columns=table_header)
    daktang = daktang.T
    daktang = daktang.iloc[4:,:]
    tg_bd = daktang.index[0]
    thoigian = datetime(bd.year,int(tg_bd[-2:]),int(tg_bd[6:8]),int(tg_bd[:2]),int(tg_bd[3:5]))
    daktang.insert(0, 'time', pd.date_range(thoigian,periods=len(daumoi.index),freq='H'))
    daktang = daktang.sort_values(by='time')
    daktang = daktang.drop(columns=[0])
    daktang.rename(columns={1:'daktang'},inplace=True)
    # print(daktang)
    
    # daknen
    select_element.select_by_value('2')
    time.sleep(2)
    element = driver.find_element(by=By.CSS_SELECTOR,value='#menu2 > div:nth-child(3) > button.btn.btn-default')
    element.click()
    time.sleep(5)
    table_element = driver.find_element(by=By.CSS_SELECTOR, value= "#data2 > table")
    table_header = []
    for th in table_element.find_elements(by=By.TAG_NAME, value="th"):
        table_header.append(th.text)

    table_data = []
    for row in table_element.find_elements(by=By.TAG_NAME,value="tr"):
        row_data = []
        for cell in row.find_elements(by=By.TAG_NAME,value="td"):
            row_data.append(cell.text)
        table_data.append(row_data)
    daknen = pd.DataFrame(table_data,columns=table_header)
    daknen = daknen.T
    daknen = daknen.iloc[4:,:]
    tg_bd = daknen.index[0]
    thoigian = datetime(bd.year,int(tg_bd[-2:]),int(tg_bd[6:8]),int(tg_bd[:2]),int(tg_bd[3:5]))
    daknen.insert(0, 'time', pd.date_range(thoigian,periods=len(daumoi.index),freq='H'))
    daknen = daknen.sort_values(by='time')
    daknen = daknen.drop(columns=[0])
    daknen.rename(columns={1:'daknen'},inplace=True)
    df = daumoi.merge(daktang,how='left',on='time')
    muadr = df.merge(daknen,how='left',on='time')

    driver.quit()
    now = datetime.now()
    kt = datetime(now.year,now.month,now.day,now.hour)
    bd = kt - timedelta(days=1)
    data = pd.DataFrame()
    data['time'] = pd.date_range(bd,kt,freq='T')
    tram = pd.read_csv('ts_id/TTB_MUA_ODA.txt')
    for item in zip(tram.Matram,tram.tentram,tram.TAB):
    # print(item[0],item[2],item[1])
        pth = 'http://113.160.225.84:2018/API_TTB/XEM/solieu.php?matram={}&ten_table={}&sophut=1&tinhtong=0&thoigianbd=%27{}%2000:00:00%27&thoigiankt=%27{}%2023:59:00%27'
        pth = pth.format(item[0],item[2],bd.strftime('%Y-%m-%d'),kt.strftime('%Y-%m-%d'))
        df = pd.read_html(pth)
        df[0].rename(columns={"thoi gian":'time','so lieu':item[1]},inplace=True)
        df = df[0].drop('Ma tram',axis=1)
        df['time'] = pd.to_datetime(df['time'])
        data = data.merge(df,how='left',on='time')
    data.set_index('time',inplace=True)
    muagio = data.rolling(60,min_periods=1).sum()
    muagio = muagio[muagio.index.minute == 0]
    epsilon = 1e-10
    muagio = muagio.applymap(lambda x: 0 if abs(x) < epsilon else x)    
    muagio =muagio.astype(float)
    df = muadr.merge(muagio,how='left',on='time')
    df = df.replace('-',np.nan)
    df.iloc[:,1:] =df.iloc[:,1:].astype(float)
    df = df[['daumoi','daknen','daktang','Son Tay']]
    return df

def vitridat():
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'
    df = pd.read_excel(pth,sheet_name='Mua')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)
    kt = now - timedelta(hours=23)
    df=df[['time','Đầu mối']]
    # print(df)
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(now.year,9,1,0), periods=len(df['time']), freq="H")
    df['time'] = dt_rang
    df = df.loc[df['time'] > kt ]
    # print(df)
    return df.index[0]
def vitridat_05day(ngay):
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'
    df = pd.read_excel(pth,sheet_name='nhiet_am',engine='openpyxl')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)

    df=df[['time','nhiet_tb']]
    # print(df)
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(now.year,8,31,0), periods=len(df['time']), freq="D")
    df['time'] = dt_rang
    df = df.loc[df['time'] >= ngay ]
    # print(df)
    return df.index[0]

def vitridat_05day_muatin(ngay):
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'
    df = pd.read_excel(pth,sheet_name='Muangay_theotin',engine='openpyxl')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)

    df=df[['time','Đầu mối']]
    # print(df)
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    df = df.loc[df['time'] >= ngay ]
    # print(df)
    return df.index[0]

def vitridat_thuyvan():
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    df = pd.read_excel(pth,sheet_name='DRHN')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,7)
    kt = now - timedelta(hours=24)
    df = df.iloc[1:,:21]
    # df=df[['time','Đầu mối']]
    
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    
    df = df.loc[df['time'] >= kt ]
    # print(df)
    return df.index[0]
def thoigianphattin():
    now = datetime.now()
    if now >= datetime(now.year,now.month,now.day,0) and  now <= datetime(now.year,now.month,now.day,3,30):
        tgpt = datetime(now.year,now.month,now.day,1,30)
    elif now >= datetime(now.year,now.month,now.day,3) and  now <= datetime(now.year,now.month,now.day,6,0):
        tgpt = datetime(now.year,now.month,now.day,4,30)    
    elif now >= datetime(now.year,now.month,now.day,6) and  now <= datetime(now.year,now.month,now.day,9,0):
        tgpt = datetime(now.year,now.month,now.day,7,30)    
    elif now >= datetime(now.year,now.month,now.day,9) and  now <= datetime(now.year,now.month,now.day,11,30):
        tgpt = datetime(now.year,now.month,now.day,10,30)        
    elif now >= datetime(now.year,now.month,now.day,12) and  now <= datetime(now.year,now.month,now.day,15,0):
        tgpt = datetime(now.year,now.month,now.day,13,30) 
    elif now >= datetime(now.year,now.month,now.day,15) and  now <= datetime(now.year,now.month,now.day,17,30):
        tgpt = datetime(now.year,now.month,now.day,16,30)     
    elif now >= datetime(now.year,now.month,now.day,18) and  now <= datetime(now.year,now.month,now.day,20,30):
        tgpt = datetime(now.year,now.month,now.day,19,30)  
    elif now >= datetime(now.year,now.month,now.day,21) and  now <= datetime(now.year,now.month,now.day,23,30):
        tgpt = datetime(now.year,now.month,now.day,22,30)  
    return tgpt

def vitridat_thuyvan_lulu():
    tgpt = thoigianphattin()
    # print(tgpt)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    df = pd.read_excel(pth,sheet_name='LULU')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)
    kt = tgpt - timedelta(hours=2.5)
    # print(kt)
    df = df.iloc[2:,:21]
    # print(df)
    
    # df=df[['time','Đầu mối']]
    
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(2023,9,1), periods=len(df['time']), freq="H")
    df['time'] = dt_rang
    # print(df)
    df = df.loc[df['time'] >= kt ]
    print(df)
    return df.index[0]
def vitridat_hn_lulu():
    tgpt = thoigianphattin()
    if tgpt.hour == 4 or tgpt.hour == 10 or tgpt.hour == 16 or tgpt.hour == 22:
        tgpt = tgpt + timedelta(hours=2.5)
    else:
        tgpt = tgpt - timedelta(hours=0.5)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    df = pd.read_excel(pth,sheet_name='DRHN')
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)
    kt = now - timedelta(hours=3)
    df = df.iloc[2:,:21]
    # print(df)
    
    # df=df[['time','Đầu mối']]
    
    # df1['time'] = pd.to_datetime(df1['time'])
    dt_rang = pd.date_range(start=datetime(now.year,8,31,13), periods=len(df['time']), freq="6H")
    df['time'] = dt_rang
    df = df.loc[df['time'] >= tgpt ]
    # print(df)
    return df.index[0]
def write_rain():
    df = loadmua()
    id = vitridat() # tim vi tri dat
    # print(df)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        df.to_excel(writer, sheet_name='Mua',startrow=id , startcol=1, header=False, index=False)
    messagebox.showinfo('Thông báo','OK!')
# print(loadmua())

def convertdocx():
    word = client.Dispatch("Word.Application")
    word.Visible = True  
    odoc = word.Documents.Open(os.getcwd() + '/SOLIEU/DAKDRINH.doc')
    odoc.SaveAs(os.getcwd() +'/SOLIEU/DAKDRINH.docx', FileFormat=16)
    odoc.Close()
    word.Quit()
def convertdocx_lulu():
    word = client.Dispatch("Word.Application")
    word.Visible = True  
    odoc = word.Documents.Open(os.getcwd() + '/SOLIEU/DAKDRINH_LULU.doc')
    odoc.SaveAs(os.getcwd() +'/SOLIEU/DAKDRINH_LULU.docx', FileFormat=16)
    odoc.Close()
    word.Quit()
def quanhe_h_w(H):
    pth  = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    df = pd.read_excel(pth,sheet_name='Z-F-W')
    df = df[['H','W']]
    df =df.iloc[3:,:]
    # df.rename(columns={'Unnamed: 2':'W'},inplace=True)
    # df = df.applymap('{0:.2f}'.format)
    # print(df)
    w = df[df['H'] == H]['W']
    return w
    

def load_sl_thuyvan():
    odoc = Document('SOLIEU/DAKDRINH.docx')
    table = odoc.tables[1]
    row_count = len(table.rows)-1
    column_count = len(table.columns)-1

    Qcmdk = []
    # q cm du kien
    for a in range(3,-1,-1):
        # print(table.cell(row_count-a, column_count).text)
        dl = table.cell(row_count-a, column_count).text
        Qcmdk.append(dl)
    Qxaongdk = [] 
    # Q xa duong ong du kien    
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-1).text
        Qxaongdk.append(dl)   
             
    Qxatrandk = []    
    # Q xa tran du kien    
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-2).text
        Qxatrandk.append(dl)  
    Qcm=[]    
    # Q chay may   
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-3).text
        Qcm.append(dl)  
    Qxaong=[]    
    # Q xa ong  
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-4).text
        Qxaong.append(dl)          
    Qxatran=[]   
    # Q xa tran
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-5).text
        Qxatran.append(dl)          
    Qden = []
    # Q ve ho
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-6).text
        Qden.append(dl)   
    H = []            
    #H ho          
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count - 7).text
        H.append(dl)

    # noi suy w
    w = quanhe_h_w(float(H[-1]))
    # thetich = pd.DataFrame(data={'thetich':w})

    data = pd.DataFrame(data={'H':H,'Q':Qden,'Qxatran':Qxatran,'Qxaong':Qxaong,'Qcm':Qcm,'Qxatrandk':Qxatrandk,'Qxaongdk':Qxaongdk,'Qcmdk':Qcmdk})
    data = data.astype(float)
    id = vitridat_thuyvan() + 2 # tim vi tri dat
    # print(id)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        data['H'].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=11, header=False, index=False)
        w.to_excel(writer, sheet_name='DRHN',startrow=id + 3 , startcol=10, header=False, index=False)
        data['Q'].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=1, header=False, index=False)
        data['Q'].to_excel(writer, sheet_name='DRHN',startrow=id + 4 , startcol=2, header=False, index=False)
        data['Qxatrandk'].to_excel(writer, sheet_name='DRHN',startrow=id + 4 , startcol=3, header=False, index=False)
        data['Qxatran'].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=4, header=False, index=False)
        data['Qxaongdk'].to_excel(writer, sheet_name='DRHN',startrow=id + 4 , startcol=5, header=False, index=False)
        data['Qxaong'].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=6, header=False, index=False)
        data['Qcmdk'].to_excel(writer, sheet_name='DRHN',startrow=id + 4 , startcol=7, header=False, index=False)
        data['Qcm'].to_excel(writer, sheet_name='DRHN',startrow=id , startcol=8, header=False, index=False)
    mo_excel()
    # messagebox.showinfo('Thông báo','OK!')
    
def load_sl_thuyvan_lulu():
    # pth =r"D:\PM_PYTHON\Dakdrinh\SOLIEU\DAKDRINH_LULU.docx"
    pth = 'SOLIEU/DAKDRINH_LULU.docx'
    odoc = Document(pth)
    table = odoc.tables[1]
    row_count = len(table.rows)-1
    column_count = len(table.columns)-1
    Qcm=[]    
    # Q chay may   
    for a in range(2,-1,-1):
        dl = table.cell(row_count-a, column_count).text
        Qcm.append(dl)  
    Qxaong=[]    
    # Q xa ong  
    for a in range(2,-1,-1):
        dl = table.cell(row_count-a, column_count-1).text
        Qxaong.append(dl)          
    Qxatran=[]   
    # Q xa tran
    for a in range(2,-1,-1):
        dl = table.cell(row_count-a, column_count-2).text
        Qxatran.append(dl)          
    Qden = []
    # Q ve ho
    for a in range(2,-1,-1):
        dl = table.cell(row_count-a, column_count-3).text
        Qden.append(dl)   
    H = []            
    #H ho          
    for a in range(2,-1,-1):
        dl = table.cell(row_count-a, column_count - 4).text
        H.append(dl)

    # print(Qden,Qcm)
    # noi suy w

    table = odoc.tables[2]
    row_count = len(table.rows)-1
    column_count = len(table.columns)-1
    Qcmdk = []
    # q cm du kien
    for a in range(3,-1,-1):
        # print(table.cell(row_count-a, column_count).text)
        dl = table.cell(row_count-a, column_count).text
        Qcmdk.append(dl)
    Qxaongdk = [] 
    # Q xa duong ong du kien    
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-1).text
        Qxaongdk.append(dl)   
             
    Qxatrandk = []    
    # Q xa tran du kien    
    for a in range(3,-1,-1):
        dl = table.cell(row_count-a, column_count-2).text
        Qxatrandk.append(dl)  

    data_td = pd.DataFrame(data={'H':H,'Q':Qden,'Qxatran':Qxatran,'Qxaong':Qxaong,'Qcm':Qcm})
    data_dk = pd.DataFrame(data={'Qxatrandk':Qxatrandk,'Qxaongdk':Qxaongdk,'Qcmdk':Qcmdk})

    data_dk = data_dk.replace('',np.nan)
    data_td = data_td.astype(float)
    data_dk = data_dk.astype(float)
    print(data_td)
    # print(data_dk)
    id_lu = vitridat_thuyvan_lulu() + 1 # tim vi tri dat
    id_hn = vitridat_hn_lulu()
    
    last_value = data_td['H'].iloc[-1]
    w = quanhe_h_w(float(last_value))
    df_h = pd.DataFrame({'H': [last_value]})
    tgpt = thoigianphattin()
    # print(tgpt)
    if tgpt.hour ==4 or tgpt.hour ==10 or tgpt.hour ==16 or tgpt.hour ==22:
        print(data_td)
        tinh_tb = data_td.mean().to_frame().transpose()
        tinh_tb.rename(columns={'Q':'qtd','Qxatran':'qxt','Qxaong':'qxacong','Qcm':'qcm'},inplace=True)
    else:
        pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
        df_tbtd = pd.read_excel(pth,sheet_name='LULU')
        df_tbtd = df_tbtd.iloc[2:,:]
        dt_rang = pd.date_range(start=datetime(2023,9,1), periods=len(df_tbtd['time']), freq="H")
        df_tbtd['time'] = dt_rang
        df_tbtd = df_tbtd[(df_tbtd['time'] < tgpt) & (df_tbtd['time'] >= tgpt - timedelta(hours=5.5))]
        print(df_tbtd)
        tinh_tb = df_tbtd.mean().to_frame().transpose()
    print(tinh_tb)
    # so lieu muc nuoc tra khuc 
    now = datetime.now()
    kt = datetime(now.year,now.month,now.day,now.hour)
    bd = kt - timedelta(days=1)
    data_mucnuoc = pd.DataFrame()
    data_mucnuoc['time'] = pd.date_range(bd,kt,freq='T')
    tram = pd.read_csv('ts_id/TTB_H_ODA.txt')
    for item in zip(tram.Matram,tram.tentram,tram.TAB):
    # print(item[0],item[2],item[1])
        pth = 'http://113.160.225.84:2018/API_TTB/XEM/solieu.php?matram={}&ten_table={}&sophut=1&tinhtong=0&thoigianbd=%27{}%2000:00:00%27&thoigiankt=%27{}%2023:59:00%27'
        pth = pth.format(item[0],item[2],bd.strftime('%Y-%m-%d'),kt.strftime('%Y-%m-%d'))
        df = pd.read_html(pth)
        df[0].rename(columns={"thoi gian":'time','so lieu':item[1]},inplace=True)
        df = df[0].drop('Ma tram',axis=1)
        df['time'] = pd.to_datetime(df['time'])
        data_mucnuoc = data_mucnuoc.merge(df,how='left',on='time')
    data_mucnuoc = data_mucnuoc[data_mucnuoc['time'].dt.minute ==0]
    data_mucnuoc.set_index('time',inplace=True)
    
    df_h_tk = pd.DataFrame({'H': [data_mucnuoc['Tra Khuc'].iloc[-1]]})
    # print(data_mucnuoc)
    
    # print(id_hn)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        data_td['H'].to_excel(writer, sheet_name='LULU',startrow=id_lu , startcol=1, header=False, index=False)
        data_td['Q'].to_excel(writer, sheet_name='LULU',startrow=id_lu , startcol=2, header=False, index=False)
        data_td['Qxatran'].to_excel(writer, sheet_name='LULU',startrow=id_lu , startcol=3, header=False, index=False)
        data_td['Qxaong'].to_excel(writer, sheet_name='LULU',startrow=id_lu , startcol=4, header=False, index=False)
        data_td['Qcm'].to_excel(writer, sheet_name='LULU',startrow=id_lu , startcol=5, header=False, index=False)
        
        data_dk['Qxatrandk'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=3, header=False, index=False)
        data_dk['Qxaongdk'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=5, header=False, index=False)
        data_dk['Qcmdk'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=7, header=False, index=False)
        df_h.to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=11, header=False, index=False)
        tinh_tb['qtd'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=1, header=False, index=False)
        tinh_tb['qxt'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=4, header=False, index=False)
        tinh_tb['qxacong'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=6, header=False, index=False)
        tinh_tb['qcm'].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=8, header=False, index=False)
        w.to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=10, header=False, index=False)
        df_h_tk.to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=20, header=False, index=False)
        # data_td['Q'].to_excel(writer, sheet_name='LULU',startrow=id + 4 , startcol=2, header=False, index=False)
        # data_dk['Qxatrandk'].to_excel(writer, sheet_name='LULU',startrow=id + 4 , startcol=3, header=False, index=False)
        # data_td['Qxatran'].to_excel(writer, sheet_name='LULU',startrow=id , startcol=4, header=False, index=False)
        # data_dk['Qxaongdk'].to_excel(writer, sheet_name='LULU',startrow=id + 4 , startcol=5, header=False, index=False)
        # data_td['Qxaong'].to_excel(writer, sheet_name='LULU',startrow=id , startcol=6, header=False, index=False)
        # data_dk['Qcmdk'].to_excel(writer, sheet_name='DRHN',startrow=id + 4 , startcol=7, header=False, index=False)
        # data_td['Qcm'].to_excel(writer, sheet_name='LULU',startrow=id , startcol=8, header=False, index=False)
        # data_td['H'].iloc[-1].to_excel(writer, sheet_name='DRHN',startrow=id_hn , startcol=11, header=False, index=False)
    mo_excel()
# load_sl_thuyvan_lulu()
def load_sl_5day():
    pth = read_txt('path_tin/DRHN.txt')
    now = datetime.now()
    ttb =[]
    tmax=[]
    tmin=[]
    amtb=[]
    ammax=[]
    ammin=[]
    tg = []
    dm = []
    dn = []
    dt =[]
    st = []
    mua_tg = []
    pp =0
    for a in range(0,10):
        pp+=1
        tttt = now - timedelta(days=a)
        tttt = datetime(tttt.year,tttt.month,tttt.day,7)
        path_tin = os.path.join(pth,'DHC_TVHN_'+ tttt.strftime('%Y%m%d') + '_0930.docx')
        odoc = Document(path_tin)
        tg.append(tttt)
        ttb.append(odoc.tables[3].cell(2,1).text) 
        tmax.append(odoc.tables[3].cell(2,2).text)
        tmin.append(odoc.tables[3].cell(2,3).text)
        amtb.append(odoc.tables[3].cell(2,4).text)
        ammax.append(odoc.tables[3].cell(2,5).text)
        ammin.append(odoc.tables[3].cell(2,6).text)
        
        for t in range(1,5):
            if t == 1:
                kk = 18
            elif t ==2:
                kk= 12
            elif t ==3:
                kk= 6
            elif t ==4:
                kk= 0
            mua_tg.append(tttt - timedelta(hours=kk))
            dm.append(odoc.tables[2].cell(1,t).text)
            dn.append(odoc.tables[2].cell(2,t).text)
            dt.append(odoc.tables[2].cell(3,t).text)
            st.append(odoc.tables[2].cell(4,t).text)
        
        
        if (tttt.strftime('%d')[-1]=='1' or tttt.strftime('%d')[-1]=='6') and ('3' not in tttt.strftime('%d') and pp>1) :
            ngay = datetime(tttt.year,tttt.month,tttt.day)
            break

    df = pd.DataFrame(data={'nhiettb':ttb,'nhietmax':tmax,'nhietmin':tmin,'amtb':amtb,'ammax':ammax,'ammin':ammin,})
    df = df.astype(float)
    df.insert(0,'time',tg)
    df = df.sort_values(by='time')
    # print(df)
    id = vitridat_05day(ngay)
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        df.iloc[:,1:].to_excel(writer, sheet_name='nhiet_am',startrow=id +1 , startcol=1, header=False, index=False)
        
    df_muatin =  pd.DataFrame(data={'dm':dm,'dn':dn,'dt':dt,'st':st})
    df_muatin.insert(0,'time',mua_tg)
    df_muatin = df_muatin.sort_values(by='time')
    df_muatin = df_muatin.replace('-',np.nan)
    df_muatin.iloc[:,1:] = df_muatin.iloc[:,1:].astype(float)
    # print(df_muatin.iloc[0]['time'])
    id = vitridat_05day_muatin(df_muatin.iloc[0]['time'])
    # print(df_muatin)
    with pd.ExcelWriter(pth,mode='a',engine='openpyxl',if_sheet_exists='overlay') as writer:   # ghi vao file co san
        df_muatin.iloc[:,1:].to_excel(writer, sheet_name='Muangay_theotin',startrow=id +1 , startcol=1, header=False, index=False)
    
    messagebox.showinfo('Thông báo','OK')
    # print(idx)
    # return df