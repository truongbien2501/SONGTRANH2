from func.sendmail import *
from datetime import datetime,timedelta
from func.Seach_file import read_txt,tim_file,read_line
from tkinter import messagebox,Tk
from docx2pdf import convert
from ftplib import FTP
import pandas as pd
import numpy as np
from pdf2image import convert_from_path
import os
import sys
from PIL import Image
import io
import  mysql.connector
from docx import Document
def creat_cxn():
    # Kết nối đến MySQL
    host = '113.160.225.84'
    user = 'qltram'
    password = 'mhq@123456'
    port = 3306
    database = 'datasolieu'
    cnx = mysql.connector.connect(host=host, user=user, password=password, port=port, database=database)
    return cnx

def query_sql(list_import,table_clounms,table_name):#TidVerticalIDVelocityForDetailMeasurement 
    if  str(list_import[2])=='nan':
        return None
    else:
        sql = 'INSERT INTO ' + table_name + '('
        gt =  " VALUES ("
        for a in range(len(table_clounms)):
            # print(table_clounms[a])
            # print(list_import[a])
            if str(list_import[a]) != 'nan':
                sql = sql + table_clounms[a]+ ','
                gt = gt + ',\'{}\''.format(list_import[a])
        sql = sql  + ')' + gt + ')'
        sql = sql.replace(',)',')')
        sql = sql.replace('(,','(')
        return sql


def insert_data(df,table_name):
    df.insert(0,'Matram','5ST')
    df['sldungduoc'] = df[df.columns[2]]
    df['maloi'] = 0
    df['chinhly'] = 0
    df = df.sort_values(by='time')
    # df =df.replace(np.nan,None)
    
    # Tạo kết nối
    cnx = creat_cxn()
    # Tạo con trỏ
    cursor = cnx.cursor(buffered=True)
    
    # Lấy danh sách các tên cột từ đối tượng con trỏ
    query = f"SELECT * FROM {table_name} LIMIT 1"
    cursor.execute(query)
    
    # Lấy danh sách các tên cột từ đối tượng con trỏ
    column_names = [column[0] for column in cursor.description]
    for index, row in df.iterrows():
        
        data = row.values.tolist()
        sql = query_sql(data,column_names,table_name)
        # print(sql)
        try:
            cursor.execute(sql)
            cnx.commit()
        except:
            pass
    cursor.close()
    cnx.close()
        
def upload_file(file_path, ftp_url, username, password):
    try:
        # Tách thành phần từ URL FTP
        url_parts = ftp_url.split("/")
        ftp_server = url_parts[2]
        # print(ftp_server)
        remote_path = "/".join(url_parts[3:]) + "/" + file_path.split('\\')[-1]
        # print(remote_path)
        # Kết nối đến server FTP
        ftp = FTP(ftp_server)
        ftp.login(username, password)
        
        # Mở file cần tải lên
        with open(file_path, 'rb') as file:
            # Tải file lên FTP
            ftp.storbinary(f'STOR {remote_path}', file)
        print("Tải lên thành công!")
    except Exception as e:
        print("Lỗi khi tải lên file:", str(e))
    finally:
        # Đóng kết nối FTP
        ftp.quit()
        
def show_message(filegui):
    result = messagebox.askquestion("Thông báo", "Bạn có muốn tiếp tục gửi file: " + filegui.split('\\')[-1])
    return result

def upload_database():
    now = datetime.now()
    now = datetime(now.year,now.month,now.day,now.hour)
    bd  = now - timedelta(days=6)
    
    pth = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    df = pd.read_excel(pth,sheet_name='DRHN')
    df.columns = df.loc[0]
    df = df.iloc[1:500,:13]
    df['Ngày giờ'] = pd.to_datetime(df['Ngày giờ'])
    df.rename(columns={'Ngày giờ':'time'},inplace=True)
    df = df[['time','Qtđ','Q xả tràn obs','Q duytridc_obs','Qcm-obs','Htđ(m)']]
    df['Qtongxa'] = df['Q xả tràn obs'] + df['Q duytridc_obs'] +df['Qcm-obs'] 
    df = df[(df['time']>=bd) & (df['time']<=now) ]

    insert_data(df[['time','Qtđ']],'ho_dakdrinh_qve')
    insert_data(df[['time','Qtongxa']],'ho_dakdrinh_qdieutiet')
    insert_data(df[['time','Htđ(m)']],'ho_dakdrinh_mucnuoc')
    messagebox.showinfo('Thông báo!','OK')
    

def bientapbantin(pth,loaibantin):
    odoc = Document(pth)    
    odoc.tables[0]._element.getparent().remove(odoc.tables[0]._element) # xoa di table so 0
    for tbl in odoc.tables[-2:]:
        cell = tbl.cell(0, 0).text
        # print(cell)
        if 'Nơi nhận' in cell:
            tbl._element.getparent().remove(tbl._element)
    
    # odoc.tables[len(odoc.tables)-2]._element.getparent().remove(odoc.tables[5]._element) # xoa di table so 5
    tin_web = 'TINMAU/{}_WEB.docx'.format(loaibantin)
    odoc.save(tin_web)
    convert(tin_web,'TINMAU/{}_WEB.pdf'.format(loaibantin))
    convert_pdf_image('TINMAU/{}_WEB.pdf'.format(loaibantin),'tin_{}.png'.format(loaibantin))
 

def upload_stream():
    filegui_datatv = read_txt('path_tin/DATA_EXCEL.txt') + '/DR_THUYVAN.xlsx'
    filegui_datakt = read_txt('path_tin/DATA_EXCEL.txt') + '/DATA_DR.xlsx'    
    upload_file(filegui_datatv,read_line('url_sever/DRHN.txt')[1],read_line('infor/dakdrinh.txt')[0],read_line('infor/dakdrinh.txt')[1]) # gui data thuy van
    upload_file(filegui_datakt,read_line('url_sever/DRHN.txt')[1],read_line('infor/dakdrinh.txt')[0],read_line('infor/dakdrinh.txt')[1]) # gui data thuy van

def convert_pdf_image(filegui,nameoutput):
    os.environ["PATH"] += os.pathsep + os.getcwd() +  r'/poppler-23.11.0\Library\bin'
    images = convert_from_path(filegui)
    # for i, image in enumerate(images):
    #     image.save('image/'+ nameoutput[:-4] + '_' + str(i) +'.png', 'PNG')
    
    dai =0
    rong =0
    for image in images:
        dai = dai + image.size[1]
        rong =image.size[0]
    new_im = Image.new('RGB', (rong, dai))
    x_offset = 0
    y_offset = 0
    for im in images:
        new_im.paste(im, (x_offset,y_offset))
        y_offset += im.size[1]
    new_im.save('image/' + nameoutput)
    
def gui_drhn(): 
    now = datetime.now()
    if now.hour > 13:
        tgpt = 'chiều'
    else:
        tgpt = 'sáng'
    # QNAM_DBKTTV_STRANH_20240414_0730.docx
    filegui = tim_file(read_txt('path_tin/DRHN.txt'),'.pdf')
    filedocx = filegui.replace('.signed','')
    filedocx = filedocx.replace('.pdf','.docx')
    result = show_message(filegui)
    if result == "yes":
        # bientapbantin(filedocx,'TVHN')
        try: 
            bientapbantin(filedocx,'TVHN')
        except:
            pass
        guimail('Bản tin {} ngày {} '.format(tgpt,now.strftime('%d/%m/%Y')),filegui,filedocx,read_txt('group_mail/songtranh.txt').replace('\n',''),read_line('infor/mail.txt')[0],read_line('infor/mail.txt')[1])
        # upload_file(filegui,read_line('url_sever/SRHN.txt')[0],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ban tin
        # upload_file('DATA\\QNAM.accdb',read_line('url_sever/SRHN.txt')[1],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ảnh
        upload_file('image\\tin_TVHN.png',read_line('url_sever/SRHN.txt')[2],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ảnh
        messagebox.showinfo("Thông báo", "Đã gửi")
    else:
        messagebox.showinfo("Thông báo", "Hủy gửi")

def gui_sthv(): 
    now = datetime.now().strftime('%Y%m%d')
    filegui = tim_file(read_txt('path_tin/DRHV.txt'),'.pdf')
    filedocx = filegui.replace('.signed','')
    filedocx = filedocx.replace('.pdf','.docx')

    result = show_message(filegui)
    if result == "yes":
        # bientapbantin(filedocx,'TVHV')
        try: 
            bientapbantin(filedocx,'TVHV')
        except:
            pass
        guimail('Bản tin hạn vừa ngày {}'.format(now.strftime('%d/%m/%Y')),filegui,filedocx,read_txt('group_mail/songtranh.txt').replace('\n',''),read_line('infor/mail.txt')[0],read_line('infor/mail.txt')[1])
        # upload_file(filegui,read_line('url_sever/SRHN.txt')[0],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ban tin
        upload_file('image\\tin_TVHV.png',read_line('url_sever/SRHN.txt')[2],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ảnh
        messagebox.showinfo("Thông báo", "Đã gửi")
    else:
        messagebox.showinfo("Thông báo", "Hủy gửi")
        
def gui_sthd(): 
    now = datetime.now().strftime('%Y%m%d')
    filegui = tim_file(read_txt('path_tin/DRHD.txt'),'.pdf')
    filedocx = filegui.replace('.signed','')
    filedocx = filedocx.replace('.pdf','.docx')
    result = show_message(filegui)
    if result == "yes":
        # bientapbantin(filedocx,'TVHD')
        try: 
            bientapbantin(filedocx,'TVHD')
        except:
            pass        
        guimail('Bản tin hạn dài ngày {}'.format(now.strftime('%d/%m/%Y')),filegui,filedocx,read_txt('group_mail/songtranh.txt').replace('\n',''),read_line('infor/mail.txt')[0],read_line('infor/mail.txt')[1])
        # upload_file(filegui,read_line('url_sever/SRHN.txt')[0],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ban tin
        upload_file('image\\tin_TVHD.png',read_line('url_sever/SRHN.txt')[2],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ảnh
        messagebox.showinfo("Thông báo", "Đã gửi")
    else:
        messagebox.showinfo("Thông báo", "Hủy gửi")

def gui_lulu(): 
    now = datetime.now().strftime('%Y%m%d')
    filegui = tim_file(read_txt('path_tin/LULU.txt'),'.pdf')
    filedocx = filegui.replace('.signed','')
    filedocx = filedocx.replace('.pdf','.docx')
    result = show_message(filegui)
    if result == "yes":
        # bientapbantin(filedocx,'LULU')
        try: 
            bientapbantin(filedocx,'LULU')
        except:
            pass
        guimail('Bản tin hạn vừa ngày {} '.format(now.strftime('%d/%m/%Y')),filegui,filedocx,read_txt('group_mail/songtranh.txt').replace('\n',''),read_line('infor/mail.txt')[0],read_line('infor/mail.txt')[1])
        # upload_file(filegui,read_line('url_sever/SRHN.txt')[0],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ban tin
        upload_file('image\\tin_LULU.png',read_line('url_sever/SRHN.txt')[2],read_line('infor/songtranh.txt')[0],read_line('infor/songtranh.txt')[1]) # gui ảnh
        messagebox.showinfo("Thông báo", "Đã gửi")
    else:
        messagebox.showinfo("Thông báo", "Hủy gửi")